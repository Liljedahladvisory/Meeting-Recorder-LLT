#!/usr/bin/env python3
"""
Meeting Recorder & Notes Generator — Powered by Liljedahl Legal Tech
"""

# ── macOS: set app name + dock icon BEFORE tkinter starts ────────────────────
import sys, os, ctypes, ctypes.util
if sys.platform == "darwin":
    try:
        from Foundation import NSBundle, NSProcessInfo
        from AppKit import NSApplication, NSImage

        _APP_NAME = "Meeting Recorder LLT"

        # 1) Set process name at C level (affects menu bar)
        _libc = ctypes.cdll.LoadLibrary(ctypes.util.find_library("c"))
        _libc.setprogname(_APP_NAME.encode("utf-8"))

        # 2) Set Cocoa process name
        NSProcessInfo.processInfo().setProcessName_(_APP_NAME)

        # 3) Set bundle metadata
        _bundle = NSBundle.mainBundle()
        _info = _bundle.localizedInfoDictionary() or _bundle.infoDictionary()
        if _info:
            _info["CFBundleName"] = _APP_NAME
            _info["CFBundleDisplayName"] = _APP_NAME

        # 4) Dock icon — look for .icns next to script or inside .app/Resources
        _app = NSApplication.sharedApplication()
        _script_dir = os.path.dirname(os.path.abspath(__file__))
        for _candidate in [
            os.path.join(_script_dir, "MeetingRecorder.icns"),
            os.path.join(_script_dir, "..", "Resources", "MeetingRecorder.icns"),
            os.path.join(_script_dir, "..", "..", "Resources", "MeetingRecorder.icns"),
        ]:
            if os.path.isfile(_candidate):
                _icon = NSImage.alloc().initWithContentsOfFile_(_candidate)
                if _icon:
                    _app.setApplicationIconImage_(_icon)
                break
    except Exception:
        pass
# ─────────────────────────────────────────────────────────────────────────────

import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog, ttk
import threading
import queue
import os
import time
import wave
import tempfile
import json
import base64
import hashlib
import subprocess
from datetime import datetime, date
import numpy as np
try:
    import keyring
    _KEYRING_AVAILABLE = True
except ImportError:
    _KEYRING_AVAILABLE = False

_KEYRING_SERVICE    = "MeetingRecorder-LLT"
_KEYRING_USERNAME   = "anthropic_api_key"
_KEYRING_HF_TOKEN   = "hf_token"
# ── Platform-specific paths ───────────────────────────────────────────────────
if sys.platform == "win32":
    _APP_DATA = os.path.join(os.environ.get("APPDATA", os.path.expanduser("~")), "MeetingRecorderLLT")
else:
    _APP_DATA = os.path.expanduser("~/.meeting-recorder-llt")

CONFIG_PATH  = os.path.join(_APP_DATA, "config.json")
LICENSE_PATH = os.path.join(_APP_DATA, "license.json")

APP_VERSION = "1.0.0"

SAMPLE_RATE   = 16000
CHANNELS      = 1
CHUNK_SECONDS = 20

# ── License verification (HMAC-SHA256 — stdlib only) ─────────────────────────
import hmac as _hmac
_LICENSE_HMAC_SECRET = bytes.fromhex(
    "ac274ebf4f3a4df69ffc72888328d9a7"
    "17b67f5f593599804ce8bd01083bc150"
)

def _get_machine_id() -> str:
    """Return a stable hash of the machine's hardware UUID (cross-platform)."""
    try:
        if sys.platform == "darwin":
            out = subprocess.check_output(
                ["ioreg", "-rd1", "-c", "IOPlatformExpertDevice"],
                text=True, timeout=5,
            )
            for line in out.splitlines():
                if "IOPlatformUUID" in line:
                    uuid = line.split('"')[-2]
                    return hashlib.sha256(uuid.encode()).hexdigest()[:32]
        elif sys.platform == "win32":
            # Method 1: Windows Registry MachineGuid (stable, no subprocess)
            try:
                import winreg
                reg = winreg.OpenKey(
                    winreg.HKEY_LOCAL_MACHINE,
                    r"SOFTWARE\Microsoft\Cryptography"
                )
                machine_guid = winreg.QueryValueEx(reg, "MachineGuid")[0]
                winreg.CloseKey(reg)
                if machine_guid:
                    return hashlib.sha256(machine_guid.encode()).hexdigest()[:32]
            except Exception:
                pass
            # Method 2: MAC address fallback (Python stdlib, always works)
            import uuid as _uuid
            mac = _uuid.getnode()
            return hashlib.sha256(str(mac).encode()).hexdigest()[:32]
    except Exception:
        pass
    # Last resort: MAC address (cross-platform)
    import uuid as _uuid
    return hashlib.sha256(str(_uuid.getnode()).encode()).hexdigest()[:32]

def _verify_license(key_str: str) -> dict | None:
    """Verify a license key (HMAC-SHA256). Returns payload dict or None."""
    try:
        # Strip whitespace, newlines, prefix, and chunk separators
        raw = key_str.strip().replace("\n", "").replace("\r", "").replace(" ", "")
        if raw.upper().startswith("LLT."):
            raw = raw[4:]
        raw = raw.replace(".", "")

        # Decode base64 (add padding if needed)
        pad = 4 - len(raw) % 4
        if pad != 4:
            raw += "=" * pad
        combined = base64.urlsafe_b64decode(raw)

        # HMAC-SHA256 digest is 32 bytes — split from the end
        # Format: payload_bytes + b"|" + hmac_digest(32 bytes)
        if len(combined) < 34:  # at least 1 byte payload + | + 32 sig
            return None
        sig = combined[-32:]
        if combined[-33:-32] != b"|":
            return None
        payload_bytes = combined[:-33]

        # Verify HMAC
        expected = _hmac.new(_LICENSE_HMAC_SECRET, payload_bytes,
                             hashlib.sha256).digest()
        if not _hmac.compare_digest(sig, expected):
            return None

        payload = json.loads(payload_bytes.decode("utf-8"))
        return payload
    except Exception:
        return None

def _check_license_file() -> tuple[bool, str, dict | None]:
    """Check stored license. Returns (valid, message, payload)."""
    if not os.path.isfile(LICENSE_PATH):
        return False, "Ingen licens hittad.", None
    try:
        data = json.loads(open(LICENSE_PATH, encoding="utf-8").read())
        key_str = data.get("key", "")
        stored_machine = data.get("machine_id", "")
    except Exception:
        return False, "Kunde inte l\u00e4sa licensfil.", None

    # Verify signature
    payload = _verify_license(key_str)
    if payload is None:
        return False, "Ogiltig licensnyckel.", None

    # Check machine binding
    current_machine = _get_machine_id()
    if stored_machine and stored_machine != current_machine:
        return False, "Licensen \u00e4r registrerad p\u00e5 en annan dator.", None

    # Check expiry
    expires = payload.get("expires", "2000-01-01")
    if date.fromisoformat(expires) < date.today():
        return False, f"Licensen gick ut {expires}. Kontakta Liljedahl Legal Tech f\u00f6r f\u00f6rnyelse.", None

    return True, "ok", payload

def _activate_license(key_str: str) -> tuple[bool, str]:
    """Activate a license key on this machine."""
    payload = _verify_license(key_str)
    if payload is None:
        return False, "Ogiltig licensnyckel. Kontrollera att du kopierat hela nyckeln."

    expires = payload.get("expires", "2000-01-01")
    if date.fromisoformat(expires) < date.today():
        return False, f"Licensnyckeln har redan g\u00e5tt ut ({expires})."

    # Save with machine binding
    os.makedirs(os.path.dirname(LICENSE_PATH), exist_ok=True)
    license_data = {
        "key": key_str.strip(),
        "machine_id": _get_machine_id(),
        "activated": date.today().isoformat(),
        "company": payload.get("company", ""),
        "email": payload.get("email", ""),
        "expires": expires,
    }
    with open(LICENSE_PATH, "w", encoding="utf-8") as f:
        json.dump(license_data, f, indent=2, ensure_ascii=False)

    return True, f"Licensen \u00e4r aktiverad! G\u00e4ller till {expires}."

# ── Colour palette ───────────────────────────────────────────────────────────
BG      = "#0E0D0C"
BG2     = "#161412"
BG3     = "#1E1B18"
BG4     = "#272320"
BORDER  = "#332E28"
BORDER2 = "#4A4238"
FG      = "#F2EEE8"   # primary text — warm white
FG2     = "#C8B89A"   # secondary text — field labels, section headers
FG3     = "#E0D4C0"   # emphasized text — transcript content
FG_DIM  = "#9A8A72"   # tertiary — hints, "Powered by", tab inactive
ACCENT  = "#E07820"   # orange brand accent
ACCENT2 = "#C05E0A"   # darker orange for hover/active
RED     = "#D95050"
GREEN   = "#4AB870"

# ── Fonts ────────────────────────────────────────────────────────────────────
FONT_LOGO1   = ("Helvetica Neue", 14, "bold")
FONT_LOGO2   = ("Helvetica Neue", 14)
FONT_POWERED = ("Helvetica Neue", 9, "italic")
FONT_SECTION = ("Helvetica Neue", 8, "bold")
FONT_H       = ("Helvetica Neue", 11, "bold")
FONT_B       = ("Helvetica Neue", 11)
FONT_S       = ("Helvetica Neue", 10)
FONT_XS      = ("Helvetica Neue", 9)
FONT_M       = ("Menlo", 10)
FONT_MS      = ("Menlo", 9)
FONT_TIMER   = ("Menlo", 12)
FONT_GEAR    = ("Helvetica Neue", 18)


# ── Language / i18n ───────────────────────────────────────────────────────────
_LANG = "sv"  # default, overridden from config on startup

_STRINGS = {
    "sv": {
        # Config section
        "KONFIGURATION": "KONFIGURATION",
        "API-nyckel": "API-nyckel",
        "visa": "visa",
        "dölj": "dölj",
        "Möte / titel": "Möte / titel",
        "Deltagare": "Deltagare",
        "namn, roll — kommaseparerade": "namn, roll — kommaseparerade",
        "Språk": "Språk",
        "Transkription": "Transkription",
        "small_hint": "snabb",
        "medium_hint": "balanserad",
        "large_hint": "bäst kvalitet",
        "Exportformat": "Exportformat",
        # Microphone section
        "MIKROFON": "MIKROFON",
        "Enhet": "Enhet",
        # Buttons
        "Starta inspelning": "Starta inspelning",
        "Avsluta möte": "Avsluta möte",
        "Generera anteckningar": "Generera anteckningar",
        "Spara": "Spara",
        "Stäng": "Stäng",
        "Avbryt": "Avbryt",
        "Aktivera": "Aktivera",
        "Avsluta": "Avsluta",
        "Registrera & starta": "Registrera & starta",
        "Jag har redan en nyckel": "Jag har redan en nyckel",
        "Ange ny licensnyckel": "Ange ny licensnyckel",
        "Spara inställningar": "Spara inställningar",
        # Tabs
        "Transkript": "Transkript",
        "Mötesanteckningar": "Mötesanteckningar",
        "Log": "Log",
        # Status
        "Redo.": "Redo.",
        "Transkriberar...": "Transkriberar...",
        "Genererar anteckningar...": "Genererar anteckningar...",
        "Inspelning pågår": "Inspelning pågår",
        "Klar": "Klar",
        "Sparad": "Sparad",
        # Settings dialog
        "Inställningar": "Inställningar",
        "Ändra namn eller företagsnamn:": "Ändra namn eller företagsnamn:",
        "Ditt namn": "Ditt namn",
        "Företagsnamn": "Företagsnamn",
        "Spara anteckningar i": "Spara anteckningar i",
        "Välj mapp": "Välj mapp",
        "Format": "Format",
        "Anthropic API-nyckel": "Anthropic API-nyckel",
        "HF-token (talarseparation)": "HF-token (talarseparation)",
        "hf_token_hint": "Valfritt — krävs för att identifiera vem som talar",
        "Välkommen till Meeting Recorder": "Välkommen till Meeting Recorder",
        "setup_message": "Ange ditt namn eller företagsnamn.\nDet används i mötesanteckningar och exporterade filer.",
        # Registration form
        "reg_title": "Meeting Recorder LLT — Registrering",
        "Registrera dig för att aktivera din licens:": "Registrera dig för att aktivera din licens:",
        "Namn *": "Namn *",
        "Adress *": "Adress *",
        "Organisationsnummer": "Organisationsnummer",
        "E-postadress *": "E-postadress *",
        "Registrerar...": "Registrerar...",
        "Välkommen": "Välkommen",
        "reg_fill_name": "Fyll i ditt namn.",
        "reg_fill_address": "Fyll i din adress.",
        "reg_fill_email": "Ange en giltig e-postadress.",
        "reg_cannot_create": "Kunde inte skapa licens: ",
        "reg_welcome": "Välkommen, {name}! Din licens är aktiverad.",
        # Activation
        "act_title": "Meeting Recorder LLT — Aktivering",
        "Ange din licensnyckel för att aktivera appen:": "Ange din licensnyckel för att aktivera appen:",
        "Klistra in din licensnyckel ovan.": "Klistra in din licensnyckel ovan.",
        # Expired
        "exp_title": "Meeting Recorder LLT — Licens",
        "Kontakta Liljedahl Legal Tech för förnyelse:": "Kontakta Liljedahl Legal Tech för att förnya din licens:\nsvante@liljedahladvisory.se",
        "revoked_msg": "Din licens har spärrats. Kontakta Liljedahl Legal Tech.",
        # Error messages
        "API-nyckel saknas": "API-nyckel saknas",
        "Fyll i Anthropic API-nyckel.": "Fyll i Anthropic API-nyckel.",
        "Ingen inspelning att transkribera.": "Ingen inspelning att transkribera.",
        "Ingen mikrofon": "Ingen mikrofon",
        "Välj en mikrofonenhet.": "Välj en mikrofonenhet.",
        # About dialog
        "Om Meeting Recorder LLT": "Om Meeting Recorder LLT",
        "Utvecklad av Liljedahl Advisory AB": "Utvecklad av Liljedahl Advisory AB",
        "copyright": "© 2025–2026 Liljedahl Advisory AB.\nAlla rättigheter förbehållna.",
        # Help dialog
        "Hjälp & FAQ": "Hjälp & FAQ",
        "help_title_line": "Hjälp & FAQ\n",
        "help_powered_line": "Meeting Recorder  ·  Powered by Liljedahl Legal Tech\n",
        "sec_kom_igang": "Kom igång",
        "kom_igang_para": ("Appen spelar in ditt möte via mikrofon, transkriberar det med "
                           "Whisper och genererar strukturerade mötesanteckningar med Claude AI."),
        "step_1_igang": "Ange ditt namn eller företagsnamn i inställningarna (⚙).",
        "step_2_igang": "Fyll i din Anthropic API-nyckel (se nedan hur du skaffar en).",
        "step_3_igang": "Ange mötets titel och deltagare.",
        "step_4_igang": "Klicka Starta inspelning när mötet börjar.",
        "step_5_igang": "Klicka Avsluta möte när mötet är klart.",
        "step_6_igang": "Klicka Generera anteckningar och välj sedan Spara.",
        "sec_api_key": "Hur skaffar jag en API-nyckel?",
        "api_key_para": ("API-nyckeln låter appen kommunicera med Claude AI. "
                         "Du betalar per användning direkt till Anthropic — inte till Liljedahl Legal Tech."),
        "step_1_api": "Gå till console.anthropic.com och skapa ett konto.",
        "step_2_api": "Klicka på 'API Keys' i menyn till vänster.",
        "step_3_api": "Klicka 'Create Key', ge den ett namn (t.ex. 'Meeting Recorder').",
        "step_4_api": "Kopiera nyckeln — den börjar med 'sk-ant-'.",
        "step_5_api": "Klistra in den i API-nyckelfältet i appen. Den sparas automatiskt.",
        "tip_api": "💡 Nyckeln sparas säkert i macOS Nyckelring och behöver bara anges en gång.",
        "sec_whisper": "Vilken transkriptionsmodell ska jag välja?",
        "step_small_whisper": "Snabbast, bra för korta möten och tydligt tal.",
        "step_medium_whisper": "Rekommenderas för de flesta möten (standard).",
        "step_large_whisper": "Bäst kvalitet, tar längre tid — lämplig för komplexa möten.",
        "sec_export": "Exportformat",
        "step_md_export": "Markdown — öppnas i Obsidian, Notion eller valfri textredigerare.",
        "step_docx_export": "Word — öppnas direkt i Microsoft Word eller Pages.",
        "step_pdf_export": "PDF — lämplig för att skicka anteckningar till kunder eller kollegor.",
        "sec_faq": "Vanliga frågor",
        "faq_q1": "Varför kan jag inte generera anteckningar direkt efter mötet?",
        "faq_a1": ("När du avslutar mötet fortsätter appen transkribera kvarvarande ljud i "
                   "bakgrunden. En statusrad visar hur långt transkriptionen kommit. Knappen "
                   "'Generera anteckningar' aktiveras automatiskt när allt är klart — "
                   "vänta bara tills statusraden försvinner. Hur lång tid det tar beror på "
                   "mötets längd och vald Whisper-modell."),
        "faq_q2": "Varför tar transkriptionen tid?",
        "faq_a2": ("Appen transkriberar i realtid medan mötet pågår och avslutar bearbetningen "
                   "direkt efter mötet. Ju längre möte och ju större modell, desto längre tid. "
                   "Large v3 ger bäst resultat men är långsammast."),
        "faq_q3": "Sparas mina inspelningar?",
        "faq_a3": ("Nej. Ljud bearbetas lokalt och sparas aldrig permanent. Enbart "
                   "transkriptet och anteckningarna sparas när du exporterar."),
        "faq_q4": "Vad kostar det att använda Claude AI?",
        "faq_a4": ("Anthropic tar betalt per antal tokens (ungefär ord). Ett vanligt möte "
                   "på 30 min kostar typiskt under 1 kr. Se aktuella priser på anthropic.com/pricing."),
        # Menu items
        "menu_help": "Hjälp",
        "menu_help_faq": "Hjälp & FAQ…",
        "menu_about": "Om Meeting Recorder LLT…",
        "menu_arkiv": "Arkiv",
        "menu_quit": "Avsluta",
        # Save dialog
        "Spara mötesanteckningar": "Spara mötesanteckningar",
        "save_done_status": "Sparat: {path}",
        "Sparfel": "Sparfel",
        # Progress / transcription labels
        "TRANSKRIBERAR": "TRANSKRIBERAR",
        "transcription_done": "Transkription klar ✓",
        "done_status": "Klart.  Klicka 'Generera anteckningar'.",
        "generating_status": "Genererar mötesanteckningar med Claude…",
        "notes_ready_status": "Anteckningar klara.  Klicka 'Spara' för att exportera.",
        "notes_done_btn": "✓  Klara",
        "generating_btn": "Genererar…",
        # Language dialog
        "lang_dialog_title": "Meeting Recorder LLT",
        "lang_dialog_prompt": "Select language / Välj språk",
        "lang_btn_sv": "🇸🇪  Svenska",
        "lang_btn_en": "🇬🇧  English",
        # Settings language row
        "language_row_label": "Språk / Language",
        "lang_restart_msg": "Starta om appen för att tillämpa det nya språket.",
        # AI prompt language instruction
        "ai_lang_instruction": "Skriv anteckningarna på svenska.",
        # Default meeting word (used in filename/content)
        "default_meeting": "Möte",
        # Transcription chunk status (inline format strings used in code)
        "transcribing_parts_left": "Transkriberar — {n} {part_word} kvar…",
        "part_singular": "del",
        "part_plural": "delar",
        "parts_done_of": "{done} av {total} delar klara",
        "transcribing_chunk": "Transkriberar del {i}  ({p} {remaining_word})…",
        "remaining_more": "kvar",
        "remaining_one": "återstår",
        # Auto-update
        "update_available": "🔔 Version {version} tillgänglig",
        "update_btn": "Uppdatera",
        "downloading": "Laddar ner...",
        "update_ready": "✅ Starta om appen för att tillämpa",
        "update_failed": "❌ Uppdatering misslyckades",
    },
    "en": {
        # Config section
        "KONFIGURATION": "CONFIGURATION",
        "API-nyckel": "API key",
        "visa": "show",
        "dölj": "hide",
        "Möte / titel": "Meeting / title",
        "Deltagare": "Participants",
        "namn, roll — kommaseparerade": "name, role — comma separated",
        "Språk": "Language",
        "Transkription": "Transcription",
        "small_hint": "fast",
        "medium_hint": "balanced",
        "large_hint": "best quality",
        "Exportformat": "Export format",
        # Microphone section
        "MIKROFON": "MICROPHONE",
        "Enhet": "Device",
        # Buttons
        "Starta inspelning": "Start recording",
        "Avsluta möte": "End meeting",
        "Generera anteckningar": "Generate notes",
        "Spara": "Save",
        "Stäng": "Close",
        "Avbryt": "Cancel",
        "Aktivera": "Activate",
        "Avsluta": "Quit",
        "Registrera & starta": "Register & start",
        "Jag har redan en nyckel": "I already have a key",
        "Ange ny licensnyckel": "Enter new license key",
        "Spara inställningar": "Save settings",
        # Tabs
        "Transkript": "Transcript",
        "Mötesanteckningar": "Meeting notes",
        "Log": "Log",
        # Status
        "Redo.": "Ready.",
        "Transkriberar...": "Transcribing...",
        "Genererar anteckningar...": "Generating notes...",
        "Inspelning pågår": "Recording in progress",
        "Klar": "Done",
        "Sparad": "Saved",
        # Settings dialog
        "Inställningar": "Settings",
        "Ändra namn eller företagsnamn:": "Change name or company name:",
        "Ditt namn": "Your name",
        "Företagsnamn": "Company name",
        "Spara anteckningar i": "Save notes to",
        "Välj mapp": "Choose folder",
        "Format": "Format",
        "Anthropic API-nyckel": "Anthropic API key",
        "HF-token (talarseparation)": "HF token (speaker ID)",
        "hf_token_hint": "Optional — required to identify who is speaking",
        "Välkommen till Meeting Recorder": "Welcome to Meeting Recorder",
        "setup_message": "Enter your name or company name.\nIt is used in meeting notes and exported files.",
        # Registration form
        "reg_title": "Meeting Recorder LLT — Registration",
        "Registrera dig för att aktivera din licens:": "Register to activate your license:",
        "Namn *": "Name *",
        "Adress *": "Address *",
        "Organisationsnummer": "Organization number",
        "E-postadress *": "Email address *",
        "Registrerar...": "Registering...",
        "Välkommen": "Welcome",
        "reg_fill_name": "Please enter your name.",
        "reg_fill_address": "Please enter your address.",
        "reg_fill_email": "Please enter a valid email address.",
        "reg_cannot_create": "Could not create license: ",
        "reg_welcome": "Welcome, {name}! Your license is activated.",
        # Activation
        "act_title": "Meeting Recorder LLT — Activation",
        "Ange din licensnyckel för att aktivera appen:": "Enter your license key to activate the app:",
        "Klistra in din licensnyckel ovan.": "Paste your license key above.",
        # Expired
        "exp_title": "Meeting Recorder LLT — License",
        "Kontakta Liljedahl Legal Tech för förnyelse:": "Contact Liljedahl Legal Tech for renewal:\nsvante@liljedahladvisory.se",
        "revoked_msg": "Your license has been revoked. Contact Liljedahl Legal Tech.",
        # Error messages
        "API-nyckel saknas": "API key missing",
        "Fyll i Anthropic API-nyckel.": "Please enter your Anthropic API key.",
        "Ingen inspelning att transkribera.": "No recording to transcribe.",
        "Ingen mikrofon": "No microphone",
        "Välj en mikrofonenhet.": "Please select a microphone device.",
        # About dialog
        "Om Meeting Recorder LLT": "About Meeting Recorder LLT",
        "Utvecklad av Liljedahl Advisory AB": "Developed by Liljedahl Advisory AB",
        "copyright": "© 2025–2026 Liljedahl Advisory AB.\nAll rights reserved.",
        # Help dialog
        "Hjälp & FAQ": "Help & FAQ",
        "help_title_line": "Help & FAQ\n",
        "help_powered_line": "Meeting Recorder  ·  Powered by Liljedahl Legal Tech\n",
        "sec_kom_igang": "Getting started",
        "kom_igang_para": ("The app records your meeting via microphone, transcribes it with "
                           "Whisper and generates structured meeting notes with Claude AI."),
        "step_1_igang": "Enter your name or company name in settings (⚙).",
        "step_2_igang": "Enter your Anthropic API key (see below on how to get one).",
        "step_3_igang": "Enter the meeting title and participants.",
        "step_4_igang": "Click Start recording when the meeting begins.",
        "step_5_igang": "Click End meeting when the meeting is done.",
        "step_6_igang": "Click Generate notes, then Save.",
        "sec_api_key": "How do I get an API key?",
        "api_key_para": ("The API key lets the app communicate with Claude AI. "
                         "You pay per usage directly to Anthropic — not to Liljedahl Legal Tech."),
        "step_1_api": "Go to console.anthropic.com and create an account.",
        "step_2_api": "Click 'API Keys' in the left menu.",
        "step_3_api": "Click 'Create Key', give it a name (e.g. 'Meeting Recorder').",
        "step_4_api": "Copy the key — it starts with 'sk-ant-'.",
        "step_5_api": "Paste it into the API key field in the app. It is saved automatically.",
        "tip_api": "💡 The key is stored securely in macOS Keychain and only needs to be entered once.",
        "sec_whisper": "Which transcription model should I choose?",
        "step_small_whisper": "Fastest, good for short meetings and clear speech.",
        "step_medium_whisper": "Recommended for most meetings (default).",
        "step_large_whisper": "Best quality, takes longer — suitable for complex meetings.",
        "sec_export": "Export formats",
        "step_md_export": "Markdown — opens in Obsidian, Notion or any text editor.",
        "step_docx_export": "Word — opens directly in Microsoft Word or Pages.",
        "step_pdf_export": "PDF — suitable for sending notes to clients or colleagues.",
        "sec_faq": "Frequently asked questions",
        "faq_q1": "Why can't I generate notes immediately after the meeting?",
        "faq_a1": ("When you end the meeting the app continues transcribing remaining audio in "
                   "the background. A status bar shows transcription progress. The "
                   "'Generate notes' button activates automatically when everything is ready — "
                   "just wait until the status bar disappears. How long it takes depends on "
                   "meeting length and the selected Whisper model."),
        "faq_q2": "Why does transcription take time?",
        "faq_a2": ("The app transcribes in real time while the meeting is ongoing and finishes "
                   "processing right after the meeting. Longer meetings and larger models take more time. "
                   "Large v3 gives the best results but is slowest."),
        "faq_q3": "Are my recordings saved?",
        "faq_a3": ("No. Audio is processed locally and never saved permanently. Only "
                   "the transcript and notes are saved when you export."),
        "faq_q4": "What does it cost to use Claude AI?",
        "faq_a4": ("Anthropic charges per number of tokens (roughly words). A typical 30-minute "
                   "meeting costs typically less than $0.10. See current prices at anthropic.com/pricing."),
        # Menu items
        "menu_help": "Help",
        "menu_help_faq": "Help & FAQ…",
        "menu_about": "About Meeting Recorder LLT…",
        "menu_arkiv": "File",
        "menu_quit": "Quit",
        # Save dialog
        "Spara mötesanteckningar": "Save meeting notes",
        "save_done_status": "Saved: {path}",
        "Sparfel": "Save error",
        # Progress / transcription labels
        "TRANSKRIBERAR": "TRANSCRIBING",
        "transcription_done": "Transcription done ✓",
        "done_status": "Done.  Click 'Generate notes'.",
        "generating_status": "Generating meeting notes with Claude…",
        "notes_ready_status": "Notes ready.  Click 'Save' to export.",
        "notes_done_btn": "✓  Done",
        "generating_btn": "Generating…",
        # Language dialog
        "lang_dialog_title": "Meeting Recorder LLT",
        "lang_dialog_prompt": "Select language / Välj språk",
        "lang_btn_sv": "🇸🇪  Svenska",
        "lang_btn_en": "🇬🇧  English",
        # Settings language row
        "language_row_label": "Language / Språk",
        "lang_restart_msg": "Restart the app to apply the new language.",
        # AI prompt language instruction
        "ai_lang_instruction": "Write the notes in English.",
        # Default meeting word
        "default_meeting": "Meeting",
        # Transcription chunk status
        "transcribing_parts_left": "Transcribing — {n} {part_word} remaining…",
        "part_singular": "part",
        "part_plural": "parts",
        "parts_done_of": "{done} of {total} parts done",
        "transcribing_chunk": "Transcribing part {i}  ({p} {remaining_word})…",
        "remaining_more": "remaining",
        "remaining_one": "left",
        # Auto-update
        "update_available": "🔔 Version {version} available",
        "update_btn": "Update",
        "downloading": "Downloading...",
        "update_ready": "✅ Restart app to apply",
        "update_failed": "❌ Update failed",
    },
}


def T(key: str) -> str:
    """Return the translated string for *key* in the current language."""
    return _STRINGS.get(_LANG, _STRINGS["sv"]).get(key, key)


def set_lang(lang: str):
    global _LANG
    _LANG = lang if lang in _STRINGS else "sv"


# ─────────────────────────────────────────────────────────────────────────────


def load_config() -> dict:
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, encoding="utf-8") as f:
                data = json.load(f)
            # Ensure language key exists with default
            if "language" not in data:
                data["language"] = "sv"
            return data
        except Exception:
            pass
    return {"language": "sv"}


def save_config(cfg: dict):
    os.makedirs(os.path.dirname(CONFIG_PATH), exist_ok=True)
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)
    except Exception:
        pass


def _check_for_updates() -> dict | None:
    """Check GitHub Releases API for a newer version. Rate-limited to once per 24 h."""
    import urllib.request
    from datetime import timedelta

    cfg = load_config()
    last_check = cfg.get("last_update_check", "")
    if last_check:
        try:
            if datetime.now() - datetime.fromisoformat(last_check) < timedelta(hours=24):
                return None
        except Exception:
            pass

    try:
        url = "https://api.github.com/repos/Liljedahladvisory/Meeting-Recorder-LLT/releases/latest"
        req = urllib.request.Request(
            url, headers={"User-Agent": f"MeetingRecorderLLT/{APP_VERSION}"}
        )
        resp = urllib.request.urlopen(req, timeout=5)
        data = json.loads(resp.read())

        # Save last check time
        cfg["last_update_check"] = datetime.now().isoformat()
        save_config(cfg)

        latest_tag = data.get("tag_name", "").lstrip("v")
        if not latest_tag:
            return None

        # Compare semantic versions
        def ver(s):
            return [int(x) for x in s.split(".")]

        if ver(latest_tag) <= ver(APP_VERSION):
            return None

        # Look for meeting_recorder.py asset (Python-only update)
        script_url = None
        for asset in data.get("assets", []):
            if asset["name"] == "meeting_recorder.py":
                script_url = asset["browser_download_url"]
                break

        return {
            "version": latest_tag,
            "script_url": script_url,        # None = full reinstall needed
            "notes": data.get("body", "")[:200],
            "html_url": data.get("html_url", ""),
        }
    except Exception:
        return None


def _apply_python_update(script_url: str) -> bool:
    """Download new meeting_recorder.py and save to ~/.meeting-recorder-llt/meeting_recorder.py."""
    import urllib.request, shutil, tempfile
    try:
        UPDATE_PATH = os.path.join(
            os.path.expanduser("~"), ".meeting-recorder-llt", "meeting_recorder.py"
        )
        with urllib.request.urlopen(script_url, timeout=30) as resp:
            data = resp.read()
        # Write to temp first, then move (atomic)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".py") as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        shutil.move(tmp_path, UPDATE_PATH)
        return True
    except Exception:
        return False


class RoundedButton(tk.Canvas):
    """Canvas-based button with smooth rounded corners and full colour control."""

    def __init__(self, parent, text, command=None, style="solid",
                 bg=None, fg="#FFFFFF", radius=12,
                 font_spec=None, padx=28, pady=12,
                 state="normal", fixed_width=None):
        import tkinter.font as tkfont

        self._style    = style          # "solid" or "ghost"
        self._bg       = bg or ACCENT
        self._fg       = fg
        self._radius   = radius
        self._padx     = padx
        self._pady     = pady
        self._command  = command
        self._enabled  = (state == "normal")
        self._text     = text
        self._hovering = False
        self._fspec    = font_spec or ("Helvetica Neue", 13)

        weight = "bold" if len(self._fspec) > 2 and "bold" in self._fspec[2] else "normal"
        mf = tkfont.Font(family=self._fspec[0], size=self._fspec[1], weight=weight)
        th = mf.metrics("linespace")
        tw = mf.measure(text)
        self._btn_w = fixed_width or (tw + 2 * padx)
        self._btn_h = th + 2 * pady

        super().__init__(parent, width=self._btn_w, height=self._btn_h,
                         bg=BG, highlightthickness=0, bd=0)
        self._draw()
        self.bind("<Button-1>", self._on_click)
        self.bind("<Enter>",    self._on_enter)
        self.bind("<Leave>",    self._on_leave)

    # ── drawing ───────────────────────────────────────────────────────────────

    def _resolve_fill(self):
        if not self._enabled:
            return (BG3, FG_DIM) if self._style == "solid" else (BG, BORDER)
        if self._hovering:
            if self._style == "solid":
                c = self._bg.lstrip("#")
                r, g, b = int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
                darker = f"#{int(r*.82):02x}{int(g*.82):02x}{int(b*.82):02x}"
                return darker, self._fg
            else:
                return BG2, self._fg
        if self._style == "solid":
            return self._bg, self._fg
        return BG, self._fg

    def _draw(self):
        self.delete("all")
        w, h, r = self._btn_w, self._btn_h, self._radius
        fill, fg = self._resolve_fill()
        outline = fg if self._style == "ghost" else fill

        # Inset by 1 px so the outline is never clipped by the canvas edge
        m = 1
        x0, y0, x1, y1 = m, m, w - m, h - m

        # Smooth rounded-rectangle polygon (inset)
        pts = [x0+r, y0,  x1-r, y0,  x1, y0,  x1, y0+r,
               x1, y1-r,  x1, y1,  x1-r, y1,  x0+r, y1,
               x0, y1,  x0, y1-r,  x0, y0+r,  x0, y0]
        self.create_polygon(pts, smooth=True,
                            fill=fill, outline=outline, width=1)
        self.create_text(w // 2, h // 2, text=self._text,
                         font=self._fspec, fill=fg, anchor="center")

    # ── public interface ──────────────────────────────────────────────────────

    def config(self, **kw):
        import tkinter.font as tkfont
        changed = False
        if "text" in kw:
            self._text = kw.pop("text")
            changed = True
        if "state" in kw:
            self._enabled = (kw.pop("state") == "normal")
            changed = True
        if "bg" in kw:
            self._bg = kw.pop("bg")
            changed = True
        if "fg" in kw:
            self._fg = kw.pop("fg")
            changed = True
        if "cursor" in kw:
            super().config(cursor=kw.pop("cursor"))
        if kw:
            super().config(**kw)
        if changed:
            self._draw()

    def cget(self, key):
        if key == "state":  return "normal" if self._enabled else "disabled"
        if key == "text":   return self._text
        if key == "bg":     return self._bg
        return super().cget(key)

    # ── event handlers ────────────────────────────────────────────────────────

    def _on_click(self, _=None):
        if self._enabled and self._command:
            self._command()

    def _on_enter(self, _=None):
        self._hovering = True;  self._draw()

    def _on_leave(self, _=None):
        self._hovering = False; self._draw()


def list_audio_devices():
    try:
        import sounddevice as sd
        devs = []
        for i, d in enumerate(sd.query_devices()):
            if d["max_input_channels"] > 0:
                devs.append({"index": i, "name": d["name"], "channels": d["max_input_channels"]})
        return devs
    except Exception:
        return []


def _set_macos_app_name():
    try:
        from Foundation import NSProcessInfo, NSBundle
        NSProcessInfo.processInfo().setProcessName_("Meeting Recorder")
        bundle = NSBundle.mainBundle()
        for d in filter(None, [bundle.localizedInfoDictionary(),
                                bundle.infoDictionary()]):
            d["CFBundleName"]        = "Meeting Recorder"
            d["CFBundleDisplayName"] = "Meeting Recorder"
    except Exception:
        pass


class MeetingRecorder(tk.Tk):
    def __init__(self):
        _set_macos_app_name()
        super().__init__()
        self.title("Meeting Recorder")
        self.configure(bg=BG)
        self.geometry("920x880")
        self.minsize(780, 720)
        self.resizable(True, True)

        self._config   = load_config()
        self.user_name = self._config.get("user_name", "")

        self.api_key        = tk.StringVar()
        self.hf_token       = tk.StringVar()
        self.meeting_title  = tk.StringVar()
        self.participants   = tk.StringVar()
        self.language       = tk.StringVar(value="sv")
        self.mic_device_idx = tk.IntVar(value=-1)
        self.whisper_size   = tk.StringVar(value="medium")
        self.save_format    = tk.StringVar(value="md")

        self.recording              = False
        self.audio_chunks           = []
        self.transcript_parts       = []
        self.rec_thread             = None
        self.log_q                  = queue.Queue()
        self.whisper_model          = None
        self._loaded_whisper_size   = None
        self.diarization_pipeline   = None
        self.anthropic_client       = None
        self.total_seconds          = 0
        self.timer_id               = None
        self.recording_start_time   = None   # set when recording begins
        self._devices               = []

        self.transcription_queue          = queue.Queue()
        self.pending_transcriptions       = 0
        self._total_to_transcribe         = 0
        self._transcription_worker_thread = None

        self._build_ui()
        self._build_menus()
        self._load_saved_key()
        self._poll_log()
        self.after(200, self._refresh_devices)
        threading.Thread(target=self._preload_whisper, daemon=True).start()

        if not self.user_name:
            self.after(300, self._show_setup_dialog)

    # ── macOS menu bar ────────────────────────────────────────────────────────

    def _build_menus(self):
        menubar = tk.Menu(self)
        self.config(menu=menubar)

        if sys.platform == "darwin":
            # ── macOS: Apple menu (system-managed) + Help menu ──
            app_menu = tk.Menu(menubar, name="apple", tearoff=0)
            menubar.add_cascade(menu=app_menu)

            help_menu = tk.Menu(menubar, name="help", tearoff=0)
            menubar.add_cascade(label=T("menu_help"), menu=help_menu)
            help_menu.add_command(label=T("menu_help_faq"), command=self._show_help_dialog)
            help_menu.add_separator()
            help_menu.add_command(label=T("menu_about"), command=self._show_about_dialog)

            # macOS native hooks
            self.createcommand("tkAboutDialog", self._show_about_dialog)
            self.createcommand("tk::mac::ShowHelp", self._show_help_dialog)

        else:
            # ── Windows: standard menu bar ──
            app_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label=T("menu_arkiv"), menu=app_menu)
            app_menu.add_command(label=T("menu_about"), command=self._show_about_dialog)
            app_menu.add_separator()
            app_menu.add_command(label=T("menu_quit"), command=self.quit)

            help_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label=T("menu_help"), menu=help_menu)
            help_menu.add_command(label=T("menu_help_faq"), command=self._show_help_dialog)

    def _show_update_banner(self, update_info: dict):
        """Show a dismissible update notification banner at the top of the window."""
        import webbrowser

        banner = tk.Frame(self, bg=BG3, pady=6)
        # Insert at top — before the header frame
        banner.pack(fill="x", side="top", before=self._header_frame)

        inner = tk.Frame(banner, bg=BG3)
        inner.pack()

        version = update_info["version"]
        msg = T("update_available").format(version=version)
        lbl = tk.Label(inner, text="", font=FONT_S, bg=BG3, fg=FG2)

        tk.Label(inner, text=msg, font=FONT_S, bg=BG3, fg=FG).pack(side="left", padx=(0, 12))

        def do_update():
            if update_info.get("script_url"):
                btn.config(state="disabled")
                lbl.config(text=T("downloading"))

                def _dl():
                    ok = _apply_python_update(update_info["script_url"])
                    self.after(0, lambda: lbl.config(
                        text=T("update_ready") if ok else T("update_failed"),
                        fg=GREEN if ok else RED,
                    ))

                threading.Thread(target=_dl, daemon=True).start()
            else:
                webbrowser.open(
                    update_info.get("html_url",
                                    "https://github.com/Liljedahladvisory/Meeting-Recorder-LLT/releases/latest")
                )

        btn = RoundedButton(inner, text=T("update_btn"), style="primary",
                            padx=14, pady=4, command=do_update)
        btn.pack(side="left", padx=4)

        lbl.pack(side="left", padx=8)

        tk.Button(inner, text="✕", font=FONT_S, bg=BG3, fg=FG_DIM,
                  bd=0, relief="flat", cursor="hand2",
                  command=banner.destroy).pack(side="left", padx=4)

    def _show_about_dialog(self):
        dlg = tk.Toplevel(self)
        dlg.title(T("Om Meeting Recorder LLT"))
        dlg.configure(bg=BG2)
        dlg.resizable(False, False)
        dlg.grab_set()

        self.update_idletasks()
        w, h = 400, 340
        x = self.winfo_x() + (self.winfo_width()  - w) // 2
        y = self.winfo_y() + (self.winfo_height() - h) // 2
        dlg.geometry(f"{w}x{h}+{x}+{y}")

        inner = tk.Frame(dlg, bg=BG2, padx=40, pady=30)
        inner.pack(fill="both", expand=True)

        # Icon placeholder — app name large
        tk.Label(
            inner, text="Meeting Recorder LLT",
            font=("Helvetica Neue", 18, "bold"), bg=BG2, fg=FG,
        ).pack(pady=(10, 2))

        tk.Label(
            inner, text="Version 1.0.0",
            font=FONT_S, bg=BG2, fg=FG2,
        ).pack(pady=(0, 16))

        # Separator
        tk.Frame(inner, bg=BORDER, height=1).pack(fill="x", pady=8)

        tk.Label(
            inner, text="Powered by Liljedahl Legal Tech",
            font=("Helvetica Neue", 11, "italic"), bg=BG2, fg=ACCENT,
        ).pack(pady=(8, 4))

        tk.Label(
            inner, text=T("Utvecklad av Liljedahl Advisory AB"),
            font=FONT_S, bg=BG2, fg=FG2,
        ).pack(pady=(0, 4))

        tk.Label(
            inner, text="www.liljedahladvisory.com",
            font=FONT_S, bg=BG2, fg=FG_DIM,
        ).pack(pady=(0, 4))

        tk.Label(
            inner,
            text=T("copyright"),
            font=("Helvetica Neue", 9), bg=BG2, fg=FG_DIM,
            justify="center",
        ).pack(pady=(8, 16))

        # Close button
        btn_bar = tk.Frame(dlg, bg=BG2, pady=10)
        btn_bar.pack(fill="x", side="bottom")
        close = RoundedButton(
            btn_bar, text=T("Stäng"), style="primary",
            padx=28, pady=8, command=dlg.destroy,
        )
        close.pack()

    # ── UI construction ──────────────────────────────────────────────────────

    def _build_ui(self):
        self._build_header()
        self._build_config()
        self._build_audio_source()
        self._build_buttons()
        self._build_progress()
        self._build_tabs()
        self._build_status()

    def _build_header(self):
        hdr = tk.Frame(self, bg=BG, padx=32, pady=20)
        hdr.pack(fill="x")
        self._header_frame = hdr

        left = tk.Frame(hdr, bg=BG)
        left.pack(side="left", fill="y")

        name_row = tk.Frame(left, bg=BG)
        name_row.pack(anchor="w")
        self._logo_name_lbl = tk.Label(
            name_row, text=self._display_name(),
            font=FONT_LOGO1, bg=BG, fg=FG)
        self._logo_name_lbl.pack(side="left")
        tk.Label(name_row, text="  Meeting Recorder",
                 font=FONT_LOGO2, bg=BG, fg=FG2).pack(side="left")

        tk.Label(left, text="Powered by Liljedahl Legal Tech",
                 font=FONT_POWERED, bg=BG, fg=FG_DIM).pack(anchor="w", pady=(3, 0))

        right = tk.Frame(hdr, bg=BG)
        right.pack(side="right", fill="y")

        self.timer_lbl = tk.Label(right, text="00:00:00",
                                  font=FONT_TIMER, bg=BG, fg=FG_DIM)
        self.timer_lbl.pack(side="right", padx=(16, 0))

        # Gear — Label, no border
        gear = tk.Label(right, text="⚙", font=FONT_GEAR,
                        bg=BG, fg=FG_DIM, cursor="hand2")
        gear.pack(side="right", padx=(10, 0))
        gear.bind("<Enter>",    lambda _: gear.config(fg=FG))
        gear.bind("<Leave>",    lambda _: gear.config(fg=FG_DIM))
        gear.bind("<Button-1>", lambda _: self._show_settings_dialog())

        # Help button
        help_lbl = tk.Label(right, text="?", font=("Helvetica Neue", 15, "bold"),
                            bg=BG, fg=FG_DIM, cursor="hand2")
        help_lbl.pack(side="right")
        help_lbl.bind("<Enter>",    lambda _: help_lbl.config(fg=FG))
        help_lbl.bind("<Leave>",    lambda _: help_lbl.config(fg=FG_DIM))
        help_lbl.bind("<Button-1>", lambda _: self._show_help_dialog())

        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")

    def _build_config(self):
        outer = tk.Frame(self, bg=BG, padx=32, pady=20)
        outer.pack(fill="x")
        self._section_label(outer, T("KONFIGURATION"))
        card = self._card(outer)

        def field(parent, label, var, show=None, hint=None):
            row = tk.Frame(parent, bg=BG2)
            row.pack(fill="x", pady=6)
            tk.Label(row, text=label, font=FONT_S, bg=BG2, fg=FG2,
                     width=16, anchor="w").pack(side="left")
            kw = dict(textvariable=var, font=FONT_M, bg=BG3, fg=FG,
                      insertbackground=FG, relief="flat", bd=0,
                      highlightthickness=1, highlightbackground=BORDER2,
                      highlightcolor=ACCENT)
            if show:
                kw["show"] = show
            e = tk.Entry(row, **kw)
            e.pack(side="left", fill="x", expand=True, ipady=7)
            if hint:
                tk.Label(row, text=hint, font=FONT_XS,
                         bg=BG2, fg=FG_DIM).pack(side="left", padx=(10, 0))
            return e, row

        self._key_entry, r1 = field(card, T("API-nyckel"), self.api_key, show="•")
        self._show_hide_btn = tk.Button(
            r1, text=T("visa"), font=FONT_XS,
            bg=BG4, fg=FG2, relief="flat", bd=0, cursor="hand2",
            padx=10, pady=4, activebackground=BORDER2, activeforeground=FG,
            command=self._toggle_key_visibility,
        )
        self._show_hide_btn.pack(side="left", padx=(6, 0))

        self._hf_entry, r2 = field(card, T("HF-token (talarseparation)"), self.hf_token,
                                   show="•", hint=T("hf_token_hint"))
        self._show_hf_btn = tk.Button(
            r2, text=T("visa"), font=FONT_XS,
            bg=BG4, fg=FG2, relief="flat", bd=0, cursor="hand2",
            padx=10, pady=4, activebackground=BORDER2, activeforeground=FG,
            command=self._toggle_hf_visibility,
        )
        self._show_hf_btn.pack(side="left", padx=(6, 0))

        field(card, T("Möte / titel"), self.meeting_title)
        field(card, T("Deltagare"), self.participants,
              hint=T("namn, roll — kommaseparerade"))

        # Language row (transcription language, not UI language)
        lang_row = tk.Frame(card, bg=BG2)
        lang_row.pack(fill="x", pady=6)
        tk.Label(lang_row, text=T("Språk"), font=FONT_S, bg=BG2, fg=FG2,
                 width=16, anchor="w").pack(side="left")
        for val, lbl in [("sv", "Svenska"), ("en", "English")]:
            tk.Radiobutton(lang_row, text=lbl, variable=self.language, value=val,
                           font=FONT_S, bg=BG3, fg=FG, selectcolor=ACCENT,
                           activebackground=BG4, activeforeground=FG,
                           indicatoron=False, relief="solid", bd=1,
                           padx=10, pady=3, cursor="hand2").pack(side="left", padx=(0, 6))

        # Whisper model row
        model_row = tk.Frame(card, bg=BG2)
        model_row.pack(fill="x", pady=6)
        tk.Label(model_row, text=T("Transkription"), font=FONT_S, bg=BG2, fg=FG2,
                 width=16, anchor="w").pack(side="left")
        for val, lbl, hint_key in [
            ("small",    "Small",    "small_hint"),
            ("medium",   "Medium",   "medium_hint"),
            ("large-v3", "Large v3", "large_hint"),
        ]:
            grp = tk.Frame(model_row, bg=BG2)
            grp.pack(side="left", padx=(0, 6))
            tk.Radiobutton(grp, text=f"{lbl}  ({T(hint_key)})", variable=self.whisper_size, value=val,
                           font=FONT_S, bg=BG3, fg=FG, selectcolor=ACCENT,
                           activebackground=BG4, activeforeground=FG,
                           indicatoron=False, relief="solid", bd=1,
                           padx=10, pady=3, cursor="hand2").pack(side="left")

    def _toggle_key_visibility(self):
        if self._key_entry.cget("show") == "•":
            self._key_entry.config(show="")
            self._show_hide_btn.config(text=T("dölj"))
        else:
            self._key_entry.config(show="•")
            self._show_hide_btn.config(text=T("visa"))

    def _toggle_hf_visibility(self):
        if self._hf_entry.cget("show") == "•":
            self._hf_entry.config(show="")
            self._show_hf_btn.config(text=T("dölj"))
        else:
            self._hf_entry.config(show="•")
            self._show_hf_btn.config(text=T("visa"))

    def _build_audio_source(self):
        outer = tk.Frame(self, bg=BG, padx=32, pady=4)
        outer.pack(fill="x")
        self._section_label(outer, T("MIKROFON"))
        card = self._card(outer)

        tk.Label(card, text=T("Enhet"), font=FONT_S, bg=BG2, fg=FG2).pack(anchor="w", pady=(0, 6))

        combo_row = tk.Frame(card, bg=BG2)
        combo_row.pack(fill="x")

        style = ttk.Style()
        style.theme_use("default")
        style.configure("Dark.TCombobox",
                        fieldbackground=BG3, background=BG4,
                        foreground=FG, selectbackground=BG3,
                        selectforeground=FG, bordercolor=BORDER2,
                        arrowcolor=ACCENT, relief="flat", padding=8)
        style.map("Dark.TCombobox",
                  fieldbackground=[("readonly", BG3), ("disabled", BG2)],
                  selectbackground=[("readonly", BG3)],
                  foreground=[("readonly", FG), ("disabled", FG2)],
                  selectforeground=[("readonly", FG)],
                  arrowcolor=[("readonly", ACCENT)])

        self._mic_combo = ttk.Combobox(combo_row, state="readonly",
                                        font=FONT_MS, style="Dark.TCombobox")
        self._mic_combo.pack(side="left", fill="x", expand=True)
        self._mic_combo.bind("<<ComboboxSelected>>", self._on_mic_select)

        tk.Button(combo_row, text="↺", font=("Helvetica Neue", 13),
                  bg=BG3, fg=FG2, relief="flat", bd=0, cursor="hand2",
                  padx=12, pady=3, activebackground=BG4, activeforeground=FG,
                  command=self._refresh_devices).pack(side="left", padx=(8, 0))

    def _build_buttons(self):
        outer = tk.Frame(self, bg=BG, padx=32, pady=20)
        outer.pack(fill="x")
        self._btn_frame = outer   # anchor for progress bar placement

        btn_row = tk.Frame(outer, bg=BG)
        btn_row.pack(fill="x")

        self.rec_btn = RoundedButton(
            btn_row, text=f"⬤  {T('Starta inspelning')}",
            style="solid", bg=ACCENT, fg="#FFFFFF",
            font_spec=("Helvetica Neue", 13, "bold"),
            padx=28, pady=13, radius=12,
            command=self._toggle_recording,
        )
        self.rec_btn.pack(side="left", padx=(0, 10))

        self.notes_btn = RoundedButton(
            btn_row, text=f"◆  {T('Generera anteckningar')}",
            style="ghost", fg=FG_DIM,
            font_spec=("Helvetica Neue", 13),
            padx=24, pady=13, radius=12,
            state="disabled", command=self._generate_notes,
            fixed_width=240,
        )
        self.notes_btn.pack(side="left", padx=(0, 10))

        self.save_btn = RoundedButton(
            btn_row, text=f"↓  {T('Spara')}",
            style="ghost", fg=FG_DIM,
            font_spec=("Helvetica Neue", 13),
            padx=24, pady=13, radius=12,
            state="disabled", command=self._save_output,
        )
        self.save_btn.pack(side="left")

        # Format selector — own row, clearly separated
        fmt_row = tk.Frame(outer, bg=BG)
        fmt_row.pack(fill="x", pady=(12, 0))

        tk.Label(fmt_row, text=T("Exportformat"), font=FONT_XS,
                 bg=BG, fg=FG2).pack(side="left", padx=(2, 14))

        for val, lbl, desc in [
            ("md",   ".md",   "Markdown"),
            ("docx", ".docx", "Word"),
            ("pdf",  ".pdf",  "PDF"),
        ]:
            btn = tk.Radiobutton(
                fmt_row, text=f"{lbl}  {desc}",
                variable=self.save_format, value=val,
                font=FONT_S, bg=BG3, fg=FG, selectcolor=ACCENT,
                activebackground=BG4, activeforeground=FG,
                indicatoron=False, relief="solid", bd=1,
                padx=10, pady=3, cursor="hand2")
            btn.pack(side="left", padx=(0, 6))

        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")

    def _build_progress(self):
        """Orange progress bar shown while transcribing after meeting ends."""
        self._progress_frame = tk.Frame(self, bg=BG)
        # Not packed until needed — shown by _show_progress_bar()

        inner = tk.Frame(self._progress_frame, bg=BG, padx=32, pady=0)
        inner.pack(fill="x")

        lbl_row = tk.Frame(inner, bg=BG)
        lbl_row.pack(fill="x", pady=(10, 5))
        tk.Label(lbl_row, text=T("TRANSKRIBERAR"),
                 font=FONT_SECTION, bg=BG, fg=ACCENT).pack(side="left")
        self._progress_label = tk.Label(lbl_row, text="",
                 font=FONT_XS, bg=BG, fg=FG2)
        self._progress_label.pack(side="left", padx=(10, 0))

        # Track (dark background strip)
        self._progress_track = tk.Canvas(inner, bg=BG3, height=6,
                                         highlightthickness=0, bd=0)
        self._progress_track.pack(fill="x", pady=(0, 10))
        self._progress_track.bind("<Configure>", self._redraw_progress)
        self._progress_fraction = 0.0

        tk.Frame(self._progress_frame, bg=BORDER, height=1).pack(fill="x")

    def _show_progress_bar(self):
        self._progress_frame.pack(fill="x", after=self._btn_frame)
        self._set_progress(0.0, "")

    def _hide_progress_bar(self):
        self._progress_frame.pack_forget()

    def _set_progress(self, fraction: float, label: str):
        self._progress_fraction = max(0.0, min(1.0, fraction))
        self._progress_label.config(text=label)
        self._redraw_progress()

    def _redraw_progress(self, _=None):
        c = self._progress_track
        c.delete("fill")
        w = c.winfo_width()
        h = c.winfo_height()
        if w <= 1:
            return
        fill_w = int(w * self._progress_fraction)
        if fill_w > 0:
            c.create_rectangle(0, 0, fill_w, h,
                                fill=ACCENT, outline=ACCENT, tags="fill")

    def _build_tabs(self):
        tab_bar = tk.Frame(self, bg=BG, padx=32)
        tab_bar.pack(fill="x")
        self.active_tab = tk.StringVar(value="transcript")
        for label, key in [(T("Transkript"), "transcript"),
                            (T("Mötesanteckningar"), "notes"),
                            (T("Log"), "log")]:
            tk.Radiobutton(tab_bar, text=label, variable=self.active_tab, value=key,
                           font=FONT_B, bg=BG, fg=FG2, selectcolor=BG,
                           activebackground=BG, activeforeground=FG,
                           indicatoron=False, relief="flat", bd=0,
                           padx=18, pady=12, cursor="hand2",
                           command=self._switch_tab).pack(side="left")
        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")

        self.text_frame = tk.Frame(self, bg=BG, padx=32, pady=14)
        self.text_frame.pack(fill="both", expand=True)

        kw = dict(bg=BG2, fg=FG3, insertbackground=FG, relief="flat", bd=0,
                  wrap="word", state="disabled", selectbackground=BORDER2,
                  highlightthickness=1, highlightbackground=BORDER)
        self.transcript_box = scrolledtext.ScrolledText(self.text_frame, font=FONT_M, **kw)
        self.notes_box      = scrolledtext.ScrolledText(self.text_frame, font=FONT_B, **kw)
        self.log_box        = scrolledtext.ScrolledText(
            self.text_frame, font=FONT_MS, fg=FG2,
            **{k: v for k, v in kw.items() if k not in ("fg",)})
        self.transcript_box.pack(fill="both", expand=True)
        self._switch_tab()

    def _build_status(self):
        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")
        self.status_var = tk.StringVar(value=T("Redo."))
        bar = tk.Frame(self, bg=BG2)
        bar.pack(fill="x")
        tk.Label(bar, textvariable=self.status_var, font=FONT_XS,
                 bg=BG2, fg=FG2, anchor="w", padx=32, pady=8).pack(side="left")
        tk.Label(bar, text="Powered by Liljedahl Legal Tech",
                 font=FONT_POWERED, bg=BG2, fg=FG_DIM, padx=32, pady=8).pack(side="right")

    # ── UI helpers ────────────────────────────────────────────────────────────

    @staticmethod
    def _section_label(parent, text):
        tk.Label(parent, text=text, font=FONT_SECTION,
                 bg=BG, fg=FG2).pack(anchor="w", pady=(0, 10))

    @staticmethod
    def _card(parent):
        card = tk.Frame(parent, bg=BG2, padx=22, pady=18,
                        highlightthickness=1, highlightbackground=BORDER)
        card.pack(fill="x", pady=(0, 4))
        return card

    def _display_name(self) -> str:
        return self.user_name if self.user_name else "Meeting"

    def _switch_tab(self):
        for w in (self.transcript_box, self.notes_box, self.log_box):
            w.pack_forget()
        {"transcript": self.transcript_box, "notes": self.notes_box,
         "log": self.log_box}[self.active_tab.get()].pack(fill="both", expand=True)

    def _set_status(self, msg): self.status_var.set(msg)
    def _log(self, msg): self.log_q.put(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}")

    def _poll_log(self):
        while not self.log_q.empty():
            self._append(self.log_box, self.log_q.get() + "\n")
        self.after(200, self._poll_log)

    def _append(self, widget, text, color=None):
        widget.config(state="normal")
        if color:
            tag = f"c{color.replace('#','')}"
            widget.tag_config(tag, foreground=color)
            widget.insert("end", text, tag)
        else:
            widget.insert("end", text)
        widget.see("end")
        widget.config(state="disabled")

    def _clear(self, widget):
        widget.config(state="normal")
        widget.delete("1.0", "end")
        widget.config(state="disabled")

    # ── Key management ────────────────────────────────────────────────────────

    def _load_saved_key(self):
        """Load API key and HF token from keychain."""
        if _KEYRING_AVAILABLE:
            try:
                saved = keyring.get_password(_KEYRING_SERVICE, _KEYRING_USERNAME)
                if saved:
                    self.api_key.set(saved)
                    self._log("API-nyckel laddad från Keychain.")
            except Exception:
                pass
            try:
                hf = keyring.get_password(_KEYRING_SERVICE, _KEYRING_HF_TOKEN)
                if hf:
                    self.hf_token.set(hf)
                    self._log("HF-token laddad från Keychain.")
            except Exception:
                pass
        else:
            self._log("Ingen sparad API-nyckel hittades.")

    def _save_key_to_keychain(self, key: str):
        if _KEYRING_AVAILABLE:
            try:
                keyring.set_password(_KEYRING_SERVICE, _KEYRING_USERNAME, key)
            except Exception:
                pass

    def _save_hf_token_to_keychain(self, token: str):
        if _KEYRING_AVAILABLE:
            try:
                keyring.set_password(_KEYRING_SERVICE, _KEYRING_HF_TOKEN, token)
            except Exception:
                pass

    # ── First-run / settings dialogs ─────────────────────────────────────────

    def _show_setup_dialog(self):
        self._open_name_dialog(
            title=T("Välkommen till Meeting Recorder"),
            message=T("setup_message"),
            is_first_run=True,
        )

    def _show_help_dialog(self):
        dlg = tk.Toplevel(self)
        dlg.title(T("Hjälp & FAQ"))
        dlg.configure(bg=BG2)
        dlg.resizable(False, False)
        dlg.grab_set()

        self.update_idletasks()
        w, h = 580, 640
        x = self.winfo_x() + (self.winfo_width()  - w) // 2
        y = self.winfo_y() + (self.winfo_height() - h) // 2
        dlg.geometry(f"{w}x{h}+{x}+{y}")

        # ── Pure tk.Text with tags — the ONLY approach that gives native
        #    macOS trackpad scroll support. No embedded windows/frames inside.
        txt = tk.Text(
            dlg,
            bg=BG, fg=FG,
            wrap="word",
            state="disabled",
            highlightthickness=0,
            bd=0,
            padx=36,
            pady=28,
            cursor="arrow",
            spacing1=0,
            spacing2=0,
            spacing3=6,
            relief="flat",
        )
        sb = ttk.Scrollbar(dlg, orient="vertical", command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        txt.pack(side="top", fill="both", expand=True)

        # ── Tags ──────────────────────────────────────────────────────────────
        txt.tag_configure("title",
                          font=("Helvetica Neue", 15, "bold"),
                          foreground=FG,
                          spacing1=0, spacing3=4)
        txt.tag_configure("powered",
                          font=("Helvetica Neue", 10, "italic"),
                          foreground=FG_DIM,
                          spacing1=0, spacing3=18)
        txt.tag_configure("section",
                          font=("Helvetica Neue", 11, "bold"),
                          foreground=ACCENT,
                          spacing1=18, spacing3=2)
        txt.tag_configure("rule",
                          font=("Helvetica Neue", 1),
                          foreground=BORDER,
                          background=BORDER,
                          spacing1=0, spacing3=8)
        txt.tag_configure("body",
                          font=("Helvetica Neue", 12),
                          foreground=FG2,
                          spacing1=4, spacing3=0,
                          lmargin1=0, lmargin2=0)
        txt.tag_configure("step_num",
                          font=("Helvetica Neue", 11, "bold"),
                          foreground=ACCENT,
                          spacing1=5, spacing3=0)
        txt.tag_configure("step_txt",
                          font=("Helvetica Neue", 12),
                          foreground=FG2,
                          spacing1=0, spacing3=0,
                          lmargin1=22, lmargin2=22)
        txt.tag_configure("tip",
                          font=("Helvetica Neue", 11, "italic"),
                          foreground=FG_DIM,
                          spacing1=6, spacing3=0,
                          lmargin1=0, lmargin2=0)
        txt.tag_configure("faq_q",
                          font=("Helvetica Neue", 12, "bold"),
                          foreground=ACCENT,
                          spacing1=10, spacing3=2)
        txt.tag_configure("faq_a",
                          font=("Helvetica Neue", 12),
                          foreground=FG2,
                          spacing1=0, spacing3=0,
                          lmargin1=0, lmargin2=0)
        txt.tag_configure("spacer",
                          font=("Helvetica Neue", 6),
                          spacing1=0, spacing3=0)

        # ── Helper to insert content ──────────────────────────────────────────
        def ins(tag, text):
            txt.configure(state="normal")
            txt.insert("end", text, tag)
            txt.configure(state="disabled")

        def section(title):
            ins("section", title + "\n")
            ins("rule", "─" * 68 + "\n")

        def para(text):
            ins("body", text + "\n")

        def step(n, text):
            ins("step_num", f"  {n}.  ")
            ins("step_txt", text + "\n")

        def tip(text):
            ins("tip", text + "\n")

        def faq(question, answer):
            ins("faq_q", question + "\n")
            ins("faq_a", answer + "\n")

        def spacer():
            ins("spacer", "\n")

        # ── Content ───────────────────────────────────────────────────────────
        ins("title",   T("help_title_line"))
        ins("powered", T("help_powered_line"))

        section(T("sec_kom_igang"))
        para(T("kom_igang_para"))
        step(1, T("step_1_igang"))
        step(2, T("step_2_igang"))
        step(3, T("step_3_igang"))
        step(4, T("step_4_igang"))
        step(5, T("step_5_igang"))
        step(6, T("step_6_igang"))

        section(T("sec_api_key"))
        para(T("api_key_para"))
        step(1, T("step_1_api"))
        step(2, T("step_2_api"))
        step(3, T("step_3_api"))
        step(4, T("step_4_api"))
        step(5, T("step_5_api"))
        tip(T("tip_api"))

        section(T("sec_whisper"))
        step("Small",    T("step_small_whisper"))
        step("Medium",   T("step_medium_whisper"))
        step("Large v3", T("step_large_whisper"))

        section(T("sec_export"))
        step(".md",   T("step_md_export"))
        step(".docx", T("step_docx_export"))
        step(".pdf",  T("step_pdf_export"))

        section(T("sec_faq"))
        faq(T("faq_q1"), T("faq_a1"))
        spacer()
        faq(T("faq_q2"), T("faq_a2"))
        spacer()
        faq(T("faq_q3"), T("faq_a3"))
        spacer()
        faq(T("faq_q4"), T("faq_a4"))
        spacer()

        # ── Close button — lives OUTSIDE the Text widget ──────────────────────
        btn_bar = tk.Frame(dlg, bg=BG2, pady=14)
        btn_bar.pack(fill="x", side="bottom")
        close = RoundedButton(btn_bar, text=T("Stäng"), style="solid",
                              bg=ACCENT, fg="#FFFFFF",
                              font_spec=("Helvetica Neue", 12),
                              padx=28, pady=9, radius=10,
                              command=dlg.destroy)
        close.pack()

    def _show_settings_dialog(self):
        self._open_settings_dialog()

    def _open_name_dialog(self, title: str, message: str, is_first_run: bool):
        dlg = tk.Toplevel(self)
        dlg.title(title)
        dlg.configure(bg=BG)
        dlg.resizable(False, False)
        dlg.grab_set()

        self.update_idletasks()
        x = self.winfo_x() + (self.winfo_width()  - 440) // 2
        y = self.winfo_y() + (self.winfo_height() - 240) // 2
        dlg.geometry(f"440x240+{x}+{y}")

        pad = tk.Frame(dlg, bg=BG, padx=36, pady=32)
        pad.pack(fill="both", expand=True)

        tk.Label(pad, text=title, font=("Helvetica Neue", 13, "bold"),
                 bg=BG, fg=FG).pack(anchor="w")
        tk.Label(pad, text=message, font=FONT_S, bg=BG, fg=FG2,
                 wraplength=368, justify="left").pack(anchor="w", pady=(8, 18))

        entry_var = tk.StringVar(value=self.user_name)
        entry = tk.Entry(pad, textvariable=entry_var, font=FONT_M,
                         bg=BG3, fg=FG, insertbackground=FG, relief="flat", bd=0,
                         highlightthickness=1, highlightbackground=BORDER2,
                         highlightcolor=ACCENT)
        entry.pack(fill="x", ipady=9)
        entry.focus_set()
        entry.select_range(0, "end")

        btn_row = tk.Frame(pad, bg=BG)
        btn_row.pack(fill="x", pady=(18, 0))

        def save():
            name = entry_var.get().strip()
            if not name:
                entry.config(highlightbackground=RED)
                return
            self.user_name = name
            self._config["user_name"] = name
            save_config(self._config)
            self._logo_name_lbl.config(text=self._display_name())
            dlg.destroy()

        save_lbl = tk.Label(btn_row, text=T("Spara"), font=FONT_B,
                             bg=ACCENT, fg="#FFFFFF", padx=24, pady=8, cursor="hand2")
        save_lbl.pack(side="right")
        save_lbl.bind("<Button-1>", lambda _: save())
        save_lbl.bind("<Enter>",  lambda _: save_lbl.config(bg=ACCENT2))
        save_lbl.bind("<Leave>",  lambda _: save_lbl.config(bg=ACCENT))

        if not is_first_run:
            cancel_lbl = tk.Label(btn_row, text=T("Avbryt"), font=FONT_B,
                                   bg=BG3, fg=FG2, padx=24, pady=8, cursor="hand2")
            cancel_lbl.pack(side="right", padx=(0, 8))
            cancel_lbl.bind("<Button-1>", lambda _: dlg.destroy())
            cancel_lbl.bind("<Enter>", lambda _: cancel_lbl.config(bg=BG4, fg=FG))
            cancel_lbl.bind("<Leave>", lambda _: cancel_lbl.config(bg=BG3, fg=FG2))

        entry.bind("<Return>", lambda _: save())
        dlg.protocol("WM_DELETE_WINDOW", save if is_first_run else dlg.destroy)
        dlg.wait_window()

    def _open_settings_dialog(self):
        """Full settings dialog with name and UI language toggle."""
        dlg = tk.Toplevel(self)
        dlg.title(T("Inställningar"))
        dlg.configure(bg=BG)
        dlg.resizable(False, False)
        dlg.grab_set()

        self.update_idletasks()
        w, h = 460, 320
        x = self.winfo_x() + (self.winfo_width()  - w) // 2
        y = self.winfo_y() + (self.winfo_height() - h) // 2
        dlg.geometry(f"{w}x{h}+{x}+{y}")

        pad = tk.Frame(dlg, bg=BG, padx=36, pady=28)
        pad.pack(fill="both", expand=True)

        tk.Label(pad, text=T("Inställningar"), font=("Helvetica Neue", 13, "bold"),
                 bg=BG, fg=FG).pack(anchor="w")
        tk.Label(pad, text=T("Ändra namn eller företagsnamn:"), font=FONT_S, bg=BG, fg=FG2,
                 wraplength=388, justify="left").pack(anchor="w", pady=(8, 14))

        # Name entry
        entry_var = tk.StringVar(value=self.user_name)
        entry = tk.Entry(pad, textvariable=entry_var, font=FONT_M,
                         bg=BG3, fg=FG, insertbackground=FG, relief="flat", bd=0,
                         highlightthickness=1, highlightbackground=BORDER2,
                         highlightcolor=ACCENT)
        entry.pack(fill="x", ipady=9)
        entry.focus_set()
        entry.select_range(0, "end")

        # Language toggle row
        lang_frame = tk.Frame(pad, bg=BG)
        lang_frame.pack(fill="x", pady=(18, 0))

        tk.Label(lang_frame, text=T("language_row_label"), font=FONT_S, bg=BG, fg=FG2,
                 width=20, anchor="w").pack(side="left")

        ui_lang_var = tk.StringVar(value=_LANG)
        for val, lbl in [("sv", "Svenska"), ("en", "English")]:
            tk.Radiobutton(lang_frame, text=lbl, variable=ui_lang_var, value=val,
                           font=FONT_S, bg=BG3, fg=FG, selectcolor=ACCENT,
                           activebackground=BG4, activeforeground=FG,
                           indicatoron=False, relief="solid", bd=1,
                           padx=10, pady=3, cursor="hand2").pack(side="left", padx=(0, 6))

        msg_lbl = tk.Label(pad, text="", font=FONT_XS, bg=BG, fg=FG_DIM)
        msg_lbl.pack(anchor="w", pady=(10, 0))

        btn_row = tk.Frame(pad, bg=BG)
        btn_row.pack(fill="x", pady=(14, 0))

        def save():
            name = entry_var.get().strip()
            if not name:
                entry.config(highlightbackground=RED)
                return
            self.user_name = name
            self._config["user_name"] = name

            chosen_lang = ui_lang_var.get()
            lang_changed = (chosen_lang != _LANG)
            self._config["language"] = chosen_lang
            save_config(self._config)
            self._logo_name_lbl.config(text=self._display_name())
            if lang_changed:
                msg_lbl.config(text=T("lang_restart_msg"), fg=ACCENT)
                dlg.after(1800, dlg.destroy)
            else:
                dlg.destroy()

        save_lbl = tk.Label(btn_row, text=T("Spara"), font=FONT_B,
                             bg=ACCENT, fg="#FFFFFF", padx=24, pady=8, cursor="hand2")
        save_lbl.pack(side="right")
        save_lbl.bind("<Button-1>", lambda _: save())
        save_lbl.bind("<Enter>",  lambda _: save_lbl.config(bg=ACCENT2))
        save_lbl.bind("<Leave>",  lambda _: save_lbl.config(bg=ACCENT))

        cancel_lbl = tk.Label(btn_row, text=T("Avbryt"), font=FONT_B,
                               bg=BG3, fg=FG2, padx=24, pady=8, cursor="hand2")
        cancel_lbl.pack(side="right", padx=(0, 8))
        cancel_lbl.bind("<Button-1>", lambda _: dlg.destroy())
        cancel_lbl.bind("<Enter>", lambda _: cancel_lbl.config(bg=BG4, fg=FG))
        cancel_lbl.bind("<Leave>", lambda _: cancel_lbl.config(bg=BG3, fg=FG2))

        entry.bind("<Return>", lambda _: save())
        dlg.protocol("WM_DELETE_WINDOW", dlg.destroy)
        dlg.wait_window()

    # ── Device handling ───────────────────────────────────────────────────────

    def _refresh_devices(self):
        self._devices = list_audio_devices()
        if not self._devices:
            self._set_status("sounddevice ej tillgängligt")
            return
        names = [f"{d['name']}  [{d['index']}]" for d in self._devices]
        self._mic_combo["values"] = names
        if self._devices:
            self._mic_combo.current(0)
            self.mic_device_idx.set(self._devices[0]["index"])
        self._log(f"{len(self._devices)} enheter hittade.")

    def _on_mic_select(self, _=None):
        idx = self._mic_combo.current()
        if idx >= 0:
            self.mic_device_idx.set(self._devices[idx]["index"])

    # ── Whisper loading ───────────────────────────────────────────────────────

    def _preload_whisper(self):
        size = self.whisper_size.get()
        self.after(0, lambda: self._set_rec_btn_enabled(False))
        try:
            from faster_whisper import WhisperModel
            from pyannote.audio import Pipeline
            self._log(f"Förladdar Whisper {size}…")
            self.after(0, lambda: self._set_status(
                f"Laddar Whisper {size}…  (inspelningsknappen aktiveras när den är klar)"))
            self.whisper_model = WhisperModel(size, device="cpu", compute_type="int8")
            self._loaded_whisper_size = size
            self._log("Whisper redo. Laddar talarseparation…")
            hf_token = self.hf_token.get().strip() or os.environ.get("HF_TOKEN", "")
            if hf_token:
                self.diarization_pipeline = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1", token=hf_token)
                self._log("Talarseparation redo.")
            else:
                self._log("HF-token saknas — talarseparation inaktiverad.")
            self.after(0, lambda: self._set_status(
                f"Redo  —  Whisper {size}" +
                (" + talarseparation" if self.diarization_pipeline else "") + " laddat."))
        except Exception as e:
            import traceback
            self._log(f"Förladdningsfel: {e}")
            self._log(traceback.format_exc())
            self.after(0, lambda: self._set_status("Redo (whisper ej förladdad)"))
        finally:
            self.after(0, lambda: self._set_rec_btn_enabled(True))

    def _load_whisper(self):
        size = self.whisper_size.get()
        try:
            from faster_whisper import WhisperModel
            self._log(f"Laddar Whisper {size}…")
            self.whisper_model = WhisperModel(size, device="cpu", compute_type="int8")
            self._loaded_whisper_size = size
            self._log("Whisper klar.")
            self.after(0, self._do_start)
        except Exception as e:
            self._log(f"Whisper-fel: {e}")
            self.after(0, lambda: messagebox.showerror("Whisper-fel", str(e)))

    # ── Recording flow ────────────────────────────────────────────────────────

    def _set_rec_btn_enabled(self, enabled: bool):
        self.rec_btn.config(state="normal" if enabled else "disabled")

    def _toggle_recording(self):
        if not self.recording:
            self._start_recording()
        else:
            self._stop_recording()

    def _start_recording(self):
        key = self.api_key.get().strip()
        if not key:
            messagebox.showerror(T("API-nyckel saknas"), T("Fyll i Anthropic API-nyckel."))
            return
        try:
            import anthropic as ac
            self.anthropic_client = ac.Anthropic(api_key=key)
        except Exception as e:
            messagebox.showerror("Fel", str(e))
            return
        self._save_key_to_keychain(key)
        hf = self.hf_token.get().strip()
        if hf:
            self._save_hf_token_to_keychain(hf)
        if self.mic_device_idx.get() < 0:
            messagebox.showerror(T("Ingen mikrofon"), T("Välj en mikrofonenhet."))
            return
        if self.whisper_model is None or self._loaded_whisper_size != self.whisper_size.get():
            self._set_status(f"Laddar Whisper {self.whisper_size.get()}…")
            self.update()
            threading.Thread(target=self._load_whisper, daemon=True).start()
            return
        self._do_start()

    def _do_start(self):
        self.recording = True
        self.audio_chunks = []
        self.transcript_parts = []
        self.total_seconds = 0
        self.recording_start_time = datetime.now()
        self.pending_transcriptions = 0
        self._clear(self.transcript_box)
        self._clear(self.notes_box)

        self.transcription_queue = queue.Queue()
        self._transcription_worker_thread = threading.Thread(
            target=self._transcription_worker, daemon=True)
        self._transcription_worker_thread.start()

        self.rec_btn.config(text=f"■  {T('Avsluta möte')}", bg=RED, fg="#FFFFFF")
        self.notes_btn.config(state="disabled", bg=BG, fg=FG_DIM)
        self.save_btn.config(state="disabled", bg=BG, fg=FG_DIM)
        self._set_status(f"● {T('Inspelning pågår')}  —  {T('MIKROFON')}")
        self._tick()
        self.rec_thread = threading.Thread(target=self._record_loop, daemon=True)
        self.rec_thread.start()
        self._log(f"Inspelning startad — modell: {self.whisper_size.get()}")

    def _stop_recording(self):
        self.recording = False
        if self.timer_id:
            self.after_cancel(self.timer_id)
            self.timer_id = None
        self.rec_btn.config(text=f"●  {T('Starta inspelning')}", bg=FG, fg=BG)
        self._set_rec_btn_enabled(False)
        self._set_status(T("Transkriberar...") + " " + T("Inspelning pågår").lower() + "…")
        self._log("Ljud stoppat. Väntar på transkription…")
        # Snapshot total chunks in queue at this moment (remainder may add 1 more)
        self._total_to_transcribe = self.pending_transcriptions
        if self._total_to_transcribe > 0:
            self._show_progress_bar()
        threading.Thread(target=self._wait_for_transcription, daemon=True).start()

    def _wait_for_transcription(self):
        if self.rec_thread:
            self.rec_thread.join(timeout=15)
        self.transcription_queue.put(None)
        if self._transcription_worker_thread:
            self._transcription_worker_thread.join(timeout=3600)
        self.after(0, self._on_done)

    def _on_done(self):
        self._set_rec_btn_enabled(True)
        self._set_status(T("done_status"))
        if self.transcript_parts:
            self.notes_btn.config(state="normal", bg=BG, fg=ACCENT)
        self._log(T("transcription_done"))
        # Fill bar to 100% briefly, then hide
        self._set_progress(1.0, T("transcription_done"))
        self.after(1800, self._hide_progress_bar)

    # ── Audio capture ─────────────────────────────────────────────────────────

    def _record_loop(self):
        import sounddevice as sd
        chunk_frames = SAMPLE_RATE * CHUNK_SECONDS
        mic_buf = []

        def mic_cb(indata, frames, t, status):
            mic_buf.extend(indata[:, 0].tolist())

        try:
            s = sd.InputStream(device=self.mic_device_idx.get(), samplerate=SAMPLE_RATE,
                               channels=CHANNELS, dtype="int16", callback=mic_cb)
            s.start()
            try:
                while self.recording:
                    time.sleep(0.1)
                    if len(mic_buf) >= chunk_frames:
                        data = np.array(mic_buf[:chunk_frames], dtype=np.int16)
                        del mic_buf[:chunk_frames]
                        self.audio_chunks.append(data)
                        self._enqueue_transcription(data, len(self.audio_chunks))
                if len(mic_buf) > SAMPLE_RATE:
                    data = np.array(mic_buf, dtype=np.int16)
                    self.audio_chunks.append(data)
                    self._enqueue_transcription(data, len(self.audio_chunks))
            finally:
                s.stop()
                s.close()
        except Exception as e:
            self._log(f"Ljudfel: {e}")
        self._log("Ljudinspelning avslutad.")

    # ── Transcription queue ───────────────────────────────────────────────────

    def _enqueue_transcription(self, chunk, idx):
        self.pending_transcriptions += 1
        if not self.recording:
            # Remainder chunk added after stop — grow the total accordingly
            self._total_to_transcribe += 1
        self.transcription_queue.put((chunk, idx))
        self.after(0, self._update_pending_status)

    def _update_pending_status(self):
        if self.pending_transcriptions > 0 and not self.recording:
            n     = self.pending_transcriptions
            total = self._total_to_transcribe
            done  = total - n
            frac  = done / total if total > 0 else 0.0
            part_word = T("part_singular") if n == 1 else T("part_plural")
            self._set_status(T("transcribing_parts_left").format(n=n, part_word=part_word))
            self._set_progress(frac, T("parts_done_of").format(done=done, total=total))

    def _transcription_worker(self):
        while True:
            item = self.transcription_queue.get()
            if item is None:
                break
            chunk, idx = item
            self._transcribe_chunk(chunk, idx)
            self.pending_transcriptions -= 1
            self.after(0, self._update_pending_status)

    # ── Transcription + diarization ───────────────────────────────────────────

    def _transcribe_chunk(self, chunk, idx):
        n_pending = self.pending_transcriptions
        self.after(0, lambda i=idx, p=n_pending: self._set_status(
            T("transcribing_chunk").format(
                i=i, p=p,
                remaining_word=T("remaining_more") if p > 1 else T("remaining_one"))))
        self._log(f"Transkriberar del {idx} ({len(chunk)/SAMPLE_RATE:.0f}s)…")
        fname = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
                fname = f.name
            with wave.open(fname, "wb") as wf:
                wf.setnchannels(CHANNELS)
                wf.setsampwidth(2)
                wf.setframerate(SAMPLE_RATE)
                wf.writeframes(chunk.tobytes())

            segs, info = self.whisper_model.transcribe(
                fname, language=self.language.get(), task="transcribe",
                beam_size=5, vad_filter=True,
                word_timestamps=bool(self.diarization_pipeline))
            segments = list(segs)

            if self.diarization_pipeline and segments:
                text = self._diarize(fname, segments)
            else:
                text = " ".join(s.text for s in segments).strip()

            if fname and os.path.exists(fname):
                os.unlink(fname)
                fname = None

            if text:
                ts = self._fmt_time(len(self.transcript_parts) * CHUNK_SECONDS)
                self.transcript_parts.append(text)
                line = f"[{ts}]\n{text}\n\n"
                self.after(0, lambda t=line: self._append(self.transcript_box, t, FG3))
                self._log(f"Del {idx} ({info.language}): {text.replace(chr(10),' ')[:80]}…")
            else:
                self._log(f"Del {idx}: tyst.")
        except Exception as e:
            self._log(f"Transkriptionsfel del {idx}: {e}")
        finally:
            if fname and os.path.exists(fname):
                try:
                    os.unlink(fname)
                except Exception:
                    pass

    def _diarize(self, fname, segments):
        try:
            n_speakers = None
            participants_str = self.participants.get().strip()
            if participants_str:
                count = len([p for p in participants_str.split(",") if p.strip()])
                if count >= 2:
                    n_speakers = count
                    self._log(f"Diarisering med {n_speakers} kända talare.")

            diarization_kw = {"num_speakers": n_speakers} if n_speakers else {}
            diarization = self.diarization_pipeline(fname, **diarization_kw)
            turns = sorted(
                [(t.start, t.end, spk)
                 for t, _, spk in diarization.itertracks(yield_label=True)],
                key=lambda x: x[0])

            def find_speaker(t: float) -> str:
                for start, end, spk in turns:
                    if start <= t <= end:
                        return spk
                if turns:
                    return min(turns, key=lambda x: min(abs(x[0]-t), abs(x[1]-t)))[2]
                return "OKÄND"

            word_entries = []
            for seg in segments:
                if hasattr(seg, "words") and seg.words:
                    for w in seg.words:
                        word_entries.append((w.start, w.end, w.word))
                else:
                    word_entries.append((seg.start, seg.end, seg.text.strip()))

            if not word_entries:
                return " ".join(s.text for s in segments).strip()

            labeled = [(find_speaker((s+e)/2), w) for s, e, w in word_entries]

            lines, cur_spk, cur_words = [], None, []
            for spk, word in labeled:
                if spk != cur_spk:
                    if cur_words:
                        lines.append(f"**{cur_spk}:** {''.join(cur_words).strip()}")
                    cur_spk, cur_words = spk, [word]
                else:
                    cur_words.append(word)
            if cur_words:
                lines.append(f"**{cur_spk}:** {''.join(cur_words).strip()}")

            result = "\n".join(lines)
            return result if result.strip() else " ".join(w for _, _, w in word_entries)
        except Exception as e:
            self._log(f"Diariseringsfel: {e}")
            return " ".join(s.text for s in segments).strip()

    # ── Notes generation ──────────────────────────────────────────────────────

    def _generate_notes(self):
        if not self.transcript_parts:
            messagebox.showinfo("Tomt", T("Ingen inspelning att transkribera."))
            return
        self.notes_btn.config(state="disabled", bg=BG3, fg=FG_DIM, text=T("generating_btn"))
        self._set_status(T("generating_status"))
        self.active_tab.set("notes")
        self._switch_tab()
        threading.Thread(target=self._call_claude, daemon=True).start()

    def _call_claude(self):
        transcript   = "\n".join(self.transcript_parts)
        title        = self.meeting_title.get().strip() or T("default_meeting")
        participants = self.participants.get().strip()
        date_str     = (self.recording_start_time or datetime.now()).strftime("%Y-%m-%d %H:%M")
        org          = self.user_name or "organisationen"
        lang_instr   = T("ai_lang_instruction")

        if _LANG == "en":
            system = (
                f"You are an expert at writing structured and actionable meeting notes "
                f"for {org}. {lang_instr}"
            )
            prompt = (
                f"Analyse the transcript and generate meeting notes.\n\n"
                f"Meeting: {title}\nDate: {date_str}\n"
                f"{'Participants: ' + participants if participants else ''}\n\n"
                f"TRANSCRIPT:\n{transcript}\n\n"
                f"# {title}\n**Date:** {date_str}\n"
                f"{'**Participants:** ' + participants if participants else ''}\n\n"
                f"## Summary\n## Decisions\n"
                f"## Action Points\n| Action | Owner | Deadline |\n|--------|-------|----------|\n\n"
                f"## Next steps\n## Discussion summary\n\n"
                f"---\n*Generated by {org} · Powered by Liljedahl Legal Tech*"
            )
        else:
            system = (
                f"Du är en expert på att skriva strukturerade och handlingsbara mötesanteckningar "
                f"för {org}. {lang_instr}"
            )
            prompt = (
                f"Analysera transkriptet och generera mötesanteckningar.\n\n"
                f"Möte: {title}\nDatum: {date_str}\n"
                f"{'Deltagare: ' + participants if participants else ''}\n\n"
                f"TRANSKRIPT:\n{transcript}\n\n"
                f"# {title}\n**Datum:** {date_str}\n"
                f"{'**Deltagare:** ' + participants if participants else ''}\n\n"
                f"## Sammanfattning\n## Beslut\n"
                f"## Action Points\n| Åtgärd | Ansvarig | Deadline |\n|--------|----------|----------|\n\n"
                f"## Nästa steg\n## Diskussion i sammandrag\n\n"
                f"---\n*Genererat av {org} · Powered by Liljedahl Legal Tech*"
            )
        try:
            resp = self.anthropic_client.messages.create(
                model="claude-sonnet-4-20250514", max_tokens=2000,
                system=system, messages=[{"role": "user", "content": prompt}])
            notes = resp.content[0].text
            self.after(0, lambda n=notes: self._show_notes(n))
        except Exception as e:
            self._log(f"Claude API error: {e}")
            self.after(0, lambda: messagebox.showerror("API error", str(e)))
            self.after(0, lambda: self.notes_btn.config(
                state="normal", bg=BG, fg=ACCENT, text=f"◆  {T('Generera anteckningar')}"))

    def _show_notes(self, notes):
        self._clear(self.notes_box)
        self._append(self.notes_box, notes)
        self.notes_btn.config(state="normal", bg=BG, fg=GREEN, text=T("notes_done_btn"))
        self.save_btn.config(state="normal", bg=BG, fg=ACCENT)
        self._set_status(T("notes_ready_status"))
        self._log(T("notes_done_btn"))

    # ── Save ──────────────────────────────────────────────────────────────────

    def _save_output(self):
        fmt     = self.save_format.get()
        title   = self.meeting_title.get().strip().replace(" ", "_") or "mote"
        default = f"{datetime.now().strftime('%Y%m%d_%H%M')}_{title}"

        ext_map = {"md": ".md", "docx": ".docx", "pdf": ".pdf"}
        ft_map  = {
            "md":   [("Markdown", "*.md"),   ("Alla", "*.*")],
            "docx": [("Word-dokument", "*.docx"), ("Alla", "*.*")],
            "pdf":  [("PDF", "*.pdf"),        ("Alla", "*.*")],
        }
        path = filedialog.asksaveasfilename(
            defaultextension=ext_map[fmt], initialfile=default,
            filetypes=ft_map[fmt], title=T("Spara mötesanteckningar"))
        if not path:
            return

        org = self.user_name or "Meeting Recorder"
        content_md = (
            f"# {self.meeting_title.get() or T('default_meeting')}\n"
            f"*{datetime.now().strftime('%Y-%m-%d %H:%M')} · {org} · Powered by Liljedahl Legal Tech*\n\n---\n\n"
            f"{self.notes_box.get('1.0', 'end').strip()}\n\n---\n\n## Fullständigt transkript\n\n"
            + "\n".join(self.transcript_parts) + "\n"
        )
        try:
            if fmt == "md":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content_md)
            elif fmt == "docx":
                self._save_as_docx(path, content_md)
            elif fmt == "pdf":
                self._save_as_pdf(path, content_md)
            self._set_status(T("save_done_status").format(path=path))
            self._log(T("save_done_status").format(path=path))
        except Exception as e:
            messagebox.showerror(T("Sparfel"), str(e))

    def _save_as_docx(self, path, md_text):
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        table_rows = []

        def flush_table():
            if not table_rows:
                return
            data_rows = [r for r in table_rows if not all(
                c.strip().replace("-","").replace(":","") == "" for c in r)]
            if not data_rows:
                table_rows.clear()
                return
            cols = len(data_rows[0])
            t = doc.add_table(rows=len(data_rows), cols=cols)
            t.style = "Table Grid"
            for ri, row in enumerate(data_rows):
                for ci, cell in enumerate(row[:cols]):
                    t.cell(ri, ci).text = cell.strip()
            table_rows.clear()

        for line in md_text.splitlines():
            if line.strip().startswith("|"):
                table_rows.append([c for c in line.strip().strip("|").split("|")])
                continue
            flush_table()
            s = line.strip()
            if s.startswith("### "):   doc.add_heading(s[4:], level=3)
            elif s.startswith("## "): doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):  doc.add_heading(s[2:], level=1)
            elif s == "---":
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
            elif s == "":
                doc.add_paragraph()
            else:
                p = doc.add_paragraph()
                for i, part in enumerate(s.split("**")):
                    if not part: continue
                    run = p.add_run(part)
                    if i % 2 == 1: run.bold = True
                    elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                        run.text = part[1:-1]; run.italic = True
        flush_table()
        doc.save(path)

    def _save_as_pdf(self, path, md_text):
        from fpdf import FPDF

        # Margins MUST be set before add_page()
        pdf = FPDF()
        L, R, T = 22, 22, 22
        pdf.set_margins(L, T, R)
        pdf.set_auto_page_break(auto=True, margin=22)
        pdf.add_page()

        # Usable width
        pw = pdf.w - L - R

        def clean(t):
            return t.replace("**", "").replace("*", "").replace("`", "").strip()

        def is_separator_row(s):
            return all(c in "-|: " for c in s)

        def write(font_style, size, line_h, text, extra_ln=0):
            pdf.set_font("Helvetica", font_style, size)
            pdf.set_x(L)
            pdf.multi_cell(pw, line_h, text, align="L")
            if extra_ln:
                pdf.ln(extra_ln)

        def render_table(rows):
            """Render a list of cell-lists as a PDF table with wrapping text."""
            if not rows:
                return
            n_cols = max(len(r) for r in rows)
            if n_cols == 0:
                return

            PAD = 2.0   # mm horizontal padding each side
            LH  = 5.0   # base line height mm

            header = rows[0] if rows else []
            body   = rows[1:] if len(rows) > 1 else []

            # Step 1 — minimum width = header text + padding (guarantees no header wrap)
            pdf.set_font("Helvetica", "B", 9)
            min_w = []
            for ci in range(n_cols):
                h_txt = header[ci].strip() if ci < len(header) else ""
                min_w.append(pdf.get_string_width(h_txt) + 2 * PAD + 3)

            total_min = sum(min_w)
            if total_min > pw:          # edge case: scale down if headers alone exceed width
                scale = pw / total_min
                min_w = [w * scale for w in min_w]

            # Step 2 — distribute remaining space by max body-content width
            remaining = pw - sum(min_w)
            body_max = [0.0] * n_cols
            pdf.set_font("Helvetica", "", 9)
            for row in body:
                for ci in range(n_cols):
                    txt = row[ci].strip() if ci < len(row) else ""
                    body_max[ci] = max(body_max[ci], pdf.get_string_width(txt))

            total_body = sum(body_max) or 1.0
            col_widths = [
                min_w[ci] + (body_max[ci] / total_body * remaining)
                for ci in range(n_cols)
            ]

            def wrapped_lines(text, avail, style=""):
                """Count lines needed for text in avail mm width."""
                pdf.set_font("Helvetica", style, 9)
                if not text:
                    return 1
                lines, line_w = 1, 0.0
                for word in text.split():
                    ww = pdf.get_string_width(word + " ")
                    if line_w + ww > avail and line_w > 0:
                        lines += 1; line_w = ww
                    else:
                        line_w += ww
                return lines

            pdf.set_draw_color(180, 180, 180)
            pdf.set_line_width(0.3)

            for ri, row in enumerate(rows):
                is_hdr = ri == 0
                st = "B" if is_hdr else ""

                # Pre-calculate row height from tallest cell
                max_lines = max(
                    wrapped_lines(
                        (row[ci].strip() if ci < len(row) else ""),
                        col_widths[ci] - 2 * PAD, st
                    )
                    for ci in range(n_cols)
                )
                row_h = max_lines * LH + 2 * PAD + 1

                # Page-break guard
                if pdf.get_y() + row_h > pdf.h - pdf.b_margin:
                    pdf.add_page()

                y0 = pdf.get_y()

                if is_hdr:
                    pdf.line(L, y0, L + pw, y0)   # top border

                # Render each cell
                for ci, cw in enumerate(col_widths):
                    txt = row[ci].strip() if ci < len(row) else ""
                    pdf.set_font("Helvetica", st, 9)
                    pdf.set_text_color(30, 30, 30)
                    pdf.set_xy(L + sum(col_widths[:ci]) + PAD, y0 + PAD)
                    pdf.multi_cell(cw - 2 * PAD, LH, txt,
                                   align="L", new_x="RIGHT", new_y="TOP")

                pdf.set_y(y0 + row_h)   # advance to bottom of row

                # Separator after header
                if is_hdr:
                    pdf.line(L, pdf.get_y(), L + pw, pdf.get_y())

            # Bottom border
            pdf.line(L, pdf.get_y(), L + pw, pdf.get_y())
            pdf.set_text_color(0, 0, 0)
            pdf.set_draw_color(0, 0, 0)
            pdf.set_line_width(0.2)
            pdf.ln(4)

        # Collect table rows across consecutive | lines
        table_buf = []
        lines = md_text.splitlines()
        i = 0
        while i < len(lines):
            s = lines[i].strip()
            if s.startswith("|"):
                if not is_separator_row(s):
                    table_buf.append([c.strip() for c in s.strip("|").split("|")])
            else:
                if table_buf:
                    render_table(table_buf)
                    table_buf = []
                if not s:
                    pdf.ln(4)
                elif s == "---":
                    pdf.ln(2)
                    y = pdf.get_y()
                    pdf.line(L, y, pdf.w - R, y)
                    pdf.ln(4)
                elif s.startswith("# "):
                    write("B", 16, 10, clean(s[2:]), extra_ln=3)
                elif s.startswith("## "):
                    write("B", 13, 8, clean(s[3:]), extra_ln=2)
                elif s.startswith("### "):
                    write("B", 11, 7, clean(s[4:]), extra_ln=1)
                else:
                    write("", 10, 6, clean(s))
            i += 1

        if table_buf:
            render_table(table_buf)

        pdf.output(path)

    # ── Timer ─────────────────────────────────────────────────────────────────

    def _tick(self):
        if not self.recording: return
        h, r = divmod(self.total_seconds, 3600)
        m, s = divmod(r, 60)
        self.timer_lbl.config(text=f"{h:02d}:{m:02d}:{s:02d}", fg=ACCENT)
        self.total_seconds += 1
        self.timer_id = self.after(1000, self._tick)

    @staticmethod
    def _fmt_time(seconds):
        m, s = divmod(int(seconds), 60)
        h, m = divmod(m, 60)
        return f"{h:02d}:{m:02d}:{s:02d}"


# ── Webhook URL for admin notifications ──────────────────────────────────────
# Replace with your Google Apps Script web app URL after setup
_ADMIN_WEBHOOK_URL = "https://script.google.com/macros/s/AKfycbxFxWDU5khMonP3qsTtl2nG5BrFP-wWuit8kAHqQ-c2X8lJbyFr7ytuj6LtRfw0s2-f/exec"
_ADMIN_EMAIL = "svante@liljedahladvisory.se"

# ── Revocation list URL (hosted on GitHub) ───────────────────────────────────
_REVOCATION_URL = "https://raw.githubusercontent.com/Liljedahladvisory/Meeting-Recorder-LLT/main/revoked.json"

def _check_revocation(email: str) -> bool:
    """Check if the given email is revoked. Returns True if REVOKED."""
    import urllib.request
    import urllib.error
    try:
        req = urllib.request.Request(_REVOCATION_URL,
                                     headers={"User-Agent": "MeetingRecorderLLT/1.0"})
        resp = urllib.request.urlopen(req, timeout=5)
        data = json.loads(resp.read().decode("utf-8"))
        revoked_emails = [e.lower().strip() for e in data.get("revoked", [])]
        return email.lower().strip() in revoked_emails
    except Exception:
        # If we can't reach the server, allow access (graceful offline)
        return False

def _generate_trial_key(name: str, email: str, company: str = "",
                        days: int = 3) -> str:
    """Generate a short-lived trial license key (HMAC-signed)."""
    created = date.today().isoformat()
    expires = (date.today() + __import__("datetime").timedelta(days=days)).isoformat()
    payload = {
        "company": company or name,
        "created": created,
        "email": email,
        "expires": expires,
        "trial": True,
    }
    payload_json = json.dumps(payload, separators=(",", ":"), sort_keys=True)
    payload_bytes = payload_json.encode("utf-8")
    sig = _hmac.new(_LICENSE_HMAC_SECRET, payload_bytes, hashlib.sha256).digest()
    combined = payload_bytes + b"|" + sig
    key_b64 = base64.urlsafe_b64encode(combined).decode("ascii")
    chunks = [key_b64[i:i+4] for i in range(0, len(key_b64), 4)]
    return "LLT." + ".".join(chunks)


def _notify_admin(reg_data: dict):
    """Send registration data to admin via webhook (fire-and-forget)."""
    import urllib.request
    import urllib.error

    # Method 1: Webhook (if configured)
    if _ADMIN_WEBHOOK_URL:
        try:
            req = urllib.request.Request(
                _ADMIN_WEBHOOK_URL,
                data=json.dumps(reg_data).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            urllib.request.urlopen(req, timeout=10)
            return
        except Exception:
            pass

    # Method 2: Fallback — save registration to local file for admin review
    reg_log = os.path.join(_APP_DATA, "registrations.json")
    os.makedirs(os.path.dirname(reg_log), exist_ok=True)
    existing = []
    if os.path.isfile(reg_log):
        try:
            existing = json.loads(open(reg_log, encoding="utf-8").read())
        except Exception:
            pass
    existing.append(reg_data)
    with open(reg_log, "w", encoding="utf-8") as f:
        json.dump(existing, f, indent=2, ensure_ascii=False)


def _show_registration_window() -> bool:
    """Show registration form for new users. Returns True if registered."""
    root = tk.Tk()
    root.title(T("reg_title"))
    root.configure(bg=BG2)
    root.resizable(False, False)

    w, h = 540, 580
    sx = root.winfo_screenwidth()
    sy = root.winfo_screenheight()
    root.geometry(f"{w}x{h}+{(sx-w)//2}+{(sy-h)//2}")

    registered = [False]

    inner = tk.Frame(root, bg=BG2, padx=40, pady=24)
    inner.pack(fill="both", expand=True)

    # Header
    tk.Label(inner, text="Meeting Recorder LLT",
             font=("Helvetica Neue", 20, "bold"), bg=BG2, fg=FG
             ).pack(pady=(8, 2))
    tk.Label(inner, text="Powered by Liljedahl Legal Tech",
             font=("Helvetica Neue", 10, "italic"), bg=BG2, fg=ACCENT
             ).pack(pady=(0, 16))
    tk.Frame(inner, bg=BORDER, height=1).pack(fill="x", pady=4)

    tk.Label(inner, text=T("Registrera dig för att aktivera din licens:"),
             font=FONT_S, bg=BG2, fg=FG2
             ).pack(anchor="w", pady=(12, 10))

    # Form fields
    fields = {}
    form = tk.Frame(inner, bg=BG2)
    form.pack(fill="x")

    field_defs = [
        ("name",    T("Namn *"),                   True),
        ("company", T("Företagsnamn"),              False),
        ("address", T("Adress *"),                  True),
        ("org_nr",  T("Organisationsnummer"),       False),
        ("email",   T("E-postadress *"),            True),
    ]

    for key, label, required in field_defs:
        tk.Label(form, text=label, font=FONT_S, bg=BG2, fg=FG2
                 ).pack(anchor="w", pady=(6, 2))
        entry = tk.Entry(form, font=FONT_S, bg=BG, fg=FG,
                         insertbackground=FG, relief="solid", bd=1,
                         highlightbackground=BORDER, highlightcolor=ACCENT)
        entry.pack(fill="x", pady=(0, 2))
        fields[key] = entry

    msg_label = tk.Label(inner, text="", font=FONT_S, bg=BG2, fg=FG2,
                         wraplength=440, justify="left")
    msg_label.pack(fill="x", pady=(10, 8))

    def do_register():
        name = fields["name"].get().strip()
        company = fields["company"].get().strip()
        address = fields["address"].get().strip()
        org_nr = fields["org_nr"].get().strip()
        email = fields["email"].get().strip()

        # Validate required fields
        if not name:
            msg_label.config(text=T("reg_fill_name"), fg=RED)
            return
        if not address:
            msg_label.config(text=T("reg_fill_address"), fg=RED)
            return
        if not email or "@" not in email:
            msg_label.config(text=T("reg_fill_email"), fg=RED)
            return

        msg_label.config(text=T("Registrerar..."), fg=FG_DIM)
        root.update()

        # Generate 12-month license key
        license_key = _generate_trial_key(name, email, company, days=365)

        # Activate the license
        ok, activate_msg = _activate_license(license_key)
        if not ok:
            msg_label.config(text=T("reg_cannot_create") + activate_msg, fg=RED)
            return

        # Save registration data locally
        reg_data = {
            "name": name,
            "company": company,
            "address": address,
            "org_nr": org_nr,
            "email": email,
            "registered": date.today().isoformat(),
            "license_expires": (date.today() + __import__("datetime").timedelta(days=365)).isoformat(),
            "machine_id": _get_machine_id(),
        }

        # Notify admin (in background thread)
        import threading
        threading.Thread(target=_notify_admin, args=(reg_data,), daemon=True).start()

        msg_label.config(
            text="✅ " + T("reg_welcome").format(name=name),
            fg=GREEN,
        )
        registered[0] = True
        root.after(2500, root.destroy)

    btn_frame = tk.Frame(inner, bg=BG2)
    btn_frame.pack(fill="x", pady=(4, 0))

    RoundedButton(
        btn_frame, text=T("Registrera & starta"), style="primary",
        padx=24, pady=10, command=do_register,
    ).pack(side="left")

    RoundedButton(
        btn_frame, text=T("Jag har redan en nyckel"), style="ghost",
        padx=20, pady=10,
        command=lambda: [setattr(root, '_go_activate', True), root.destroy()],
    ).pack(side="right")

    root._go_activate = False
    root.mainloop()

    if getattr(root, '_go_activate', False):
        return _show_activation_window()
    return registered[0]


def _show_activation_window() -> bool:
    """Show a license activation window. Returns True if activated."""
    root = tk.Tk()
    root.title(T("act_title"))
    root.configure(bg=BG2)
    root.resizable(False, False)

    w, h = 520, 420
    sx = root.winfo_screenwidth()
    sy = root.winfo_screenheight()
    root.geometry(f"{w}x{h}+{(sx-w)//2}+{(sy-h)//2}")

    activated = [False]

    inner = tk.Frame(root, bg=BG2, padx=40, pady=30)
    inner.pack(fill="both", expand=True)

    tk.Label(inner, text="Meeting Recorder LLT",
             font=("Helvetica Neue", 20, "bold"), bg=BG2, fg=FG
             ).pack(pady=(10, 2))
    tk.Label(inner, text="Powered by Liljedahl Legal Tech",
             font=("Helvetica Neue", 10, "italic"), bg=BG2, fg=ACCENT
             ).pack(pady=(0, 20))

    tk.Frame(inner, bg=BORDER, height=1).pack(fill="x", pady=8)

    tk.Label(inner, text=T("Ange din licensnyckel för att aktivera appen:"),
             font=FONT_S, bg=BG2, fg=FG2
             ).pack(anchor="w", pady=(16, 6))

    key_entry = tk.Text(inner, height=4, width=50,
                        font=("Menlo", 10), bg=BG, fg=FG,
                        insertbackground=FG, relief="solid", bd=1,
                        highlightbackground=BORDER, highlightcolor=ACCENT,
                        wrap="char")
    key_entry.pack(fill="x", pady=(0, 8))

    msg_label = tk.Label(inner, text="", font=FONT_S, bg=BG2, fg=FG2,
                         wraplength=420, justify="left")
    msg_label.pack(fill="x", pady=(4, 12))

    def do_activate():
        key_text = key_entry.get("1.0", "end").strip()
        if not key_text:
            msg_label.config(text=T("Klistra in din licensnyckel ovan."), fg=ACCENT)
            return
        ok, msg = _activate_license(key_text)
        if ok:
            msg_label.config(text="✅ " + msg, fg=GREEN)
            activated[0] = True
            root.after(1500, root.destroy)
        else:
            msg_label.config(text="❌ " + msg, fg=RED)

    btn_frame = tk.Frame(inner, bg=BG2)
    btn_frame.pack(fill="x")

    RoundedButton(
        btn_frame, text=T("Aktivera"), style="primary",
        padx=28, pady=10, command=do_activate,
    ).pack(side="left")

    RoundedButton(
        btn_frame, text=T("Avsluta"), style="ghost",
        padx=28, pady=10, command=root.destroy,
    ).pack(side="right")

    root.mainloop()
    return activated[0]


def _show_expired_window(message: str) -> bool:
    """Show license expired window with option to re-enter key."""
    root = tk.Tk()
    root.title(T("exp_title"))
    root.configure(bg=BG2)
    root.resizable(False, False)

    w, h = 480, 280
    sx = root.winfo_screenwidth()
    sy = root.winfo_screenheight()
    root.geometry(f"{w}x{h}+{(sx-w)//2}+{(sy-h)//2}")

    reactivate = [False]

    inner = tk.Frame(root, bg=BG2, padx=40, pady=30)
    inner.pack(fill="both", expand=True)

    tk.Label(inner, text="Meeting Recorder LLT",
             font=("Helvetica Neue", 16, "bold"), bg=BG2, fg=FG
             ).pack(pady=(10, 16))

    tk.Label(inner, text=message,
             font=FONT_S, bg=BG2, fg=RED, wraplength=380, justify="center"
             ).pack(pady=(0, 8))

    tk.Label(inner, text=T("Kontakta Liljedahl Legal Tech för förnyelse:"),
             font=("Helvetica Neue", 10), bg=BG2, fg=FG_DIM, justify="center"
             ).pack(pady=(8, 20))

    btn_frame = tk.Frame(inner, bg=BG2)
    btn_frame.pack(fill="x")

    def do_reactivate():
        reactivate[0] = True
        root.destroy()

    RoundedButton(
        btn_frame, text=T("Ange ny licensnyckel"), style="primary",
        padx=24, pady=10, command=do_reactivate,
    ).pack(side="left")

    RoundedButton(
        btn_frame, text=T("Avsluta"), style="ghost",
        padx=24, pady=10, command=root.destroy,
    ).pack(side="right")

    root.mainloop()
    return reactivate[0]


def _show_language_selection_window():
    """Show a one-time language selection dialog. Returns the chosen language code."""
    root = tk.Tk()
    root.title(T("lang_dialog_title"))
    root.configure(bg=BG2)
    root.resizable(False, False)

    w, h = 380, 220
    sx = root.winfo_screenwidth()
    sy = root.winfo_screenheight()
    root.geometry(f"{w}x{h}+{(sx-w)//2}+{(sy-h)//2}")

    chosen = [_LANG]

    inner = tk.Frame(root, bg=BG2, padx=40, pady=30)
    inner.pack(fill="both", expand=True)

    tk.Label(inner, text="Meeting Recorder LLT",
             font=("Helvetica Neue", 16, "bold"), bg=BG2, fg=FG
             ).pack(pady=(0, 4))

    tk.Label(inner, text=T("lang_dialog_prompt"),
             font=("Helvetica Neue", 12), bg=BG2, fg=FG2
             ).pack(pady=(0, 20))

    btn_row = tk.Frame(inner, bg=BG2)
    btn_row.pack()

    def pick(lang):
        chosen[0] = lang
        set_lang(lang)
        root.destroy()

    RoundedButton(
        btn_row, text=T("lang_btn_sv"), style="solid", bg=ACCENT,
        padx=20, pady=10, command=lambda: pick("sv"),
    ).pack(side="left", padx=(0, 12))

    RoundedButton(
        btn_row, text=T("lang_btn_en"), style="ghost", fg=FG,
        padx=20, pady=10, command=lambda: pick("en"),
    ).pack(side="left")

    root.mainloop()
    return chosen[0]


def _main():
    import multiprocessing
    multiprocessing.freeze_support()
    multiprocessing.set_start_method("spawn", force=True)

    # ── Load language preference early ───────────────────────────────
    is_first_run = not os.path.isfile(CONFIG_PATH)
    cfg = load_config()
    set_lang(cfg.get("language", "sv"))

    # ── First-run language selection ─────────────────────────────────
    if is_first_run:
        chosen_lang = _show_language_selection_window()
        cfg["language"] = chosen_lang
        save_config(cfg)

    # ── License check ────────────────────────────────────────────────────
    valid, msg, payload = _check_license_file()

    if not valid and os.path.isfile(LICENSE_PATH):
        # License exists but expired/invalid
        if _show_expired_window(msg):
            if not _show_activation_window():
                return
            valid, msg, payload = _check_license_file()
            if not valid:
                return
        else:
            return

    if not valid and not os.path.isfile(LICENSE_PATH):
        # No license at all — show registration form (new user)
        if not _show_registration_window():
            return
        valid, msg, payload = _check_license_file()
        if not valid:
            return

    # ── Revocation check (non-blocking — allows offline use) ──────────
    if valid and payload:
        user_email = payload.get("email", "")
        if user_email and _check_revocation(user_email):
            _show_expired_window(T("revoked_msg"))
            return
    # ─────────────────────────────────────────────────────────────────────

    app = MeetingRecorder()

    # Check for updates in background (non-blocking)
    def _bg_update_check():
        result = _check_for_updates()
        if result:
            app.after(0, lambda: app._show_update_banner(result))

    threading.Thread(target=_bg_update_check, daemon=True).start()

    app.mainloop()


if __name__ == "__main__":
    _main()
